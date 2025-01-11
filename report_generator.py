# report_generator.py
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io
from docx.enum.section import WD_ORIENT

from database.db_connection import DatabaseConnection
from database.repositories.calculation_repo import CalculationResultRepository
from models.equipment import EquipmentType
from models.substance import SubstanceType


class ReportGenerator:
    """Класс для генерации отчетов в Word"""

    def __init__(self, db: DatabaseConnection):
        self.db = db
        self.calc_repo = CalculationResultRepository(db)
        self.doc = Document()

    def _sum_data_for_fn(self, data: list):
        '''
        Функция вычисления суммирования вероятностей F при которой пострадало не менее N человек
        :param data: данные вида [[3.8e-08, 1],[5.8e-08, 2],[1.1e-08, 1]..]
        :return: данные вида: {1: 0.00018, 2: 0.012, 3: 6.9008e-06, 4: 3.8e-08, 5: 7.29e-05}
        '''
        uniq = set(sorted([i[1] for i in data]))
        result = dict(zip(uniq, [0] * len(uniq)))

        for item_data in data:
            for item_uniq in uniq:
                if item_data[1] >= item_uniq:
                    result[item_uniq] = result[item_uniq] + item_data[0]

        if 0 in result:
            del result[0]  # удалить суммарную вероятность где пострадало 0 человек
        return result

    def _sum_data_for_fg(self, data: list):
        '''
        Функция вычисления суммирования вероятностей F при которой ущерб не менее G млн.руб
        :param data: данные вида [[3.8e-08, 1.2],[5.8e-08, 0.2],[1.1e-08, 12.4]..]
        :return: данные вида: {0.2: 0.00018, 1: 0.012, 3: 6.9008e-06, 5: 3.8e-08, 6.25: 7.29e-05}
        '''
        uniq = np.arange(0, max([i[1] for i in data])+max([i[1] for i in data]) / 7, max([i[1] for i in data]) / 7)

        result = dict(zip(uniq, [0] * len(uniq)))

        for item_data in data:
            for item_uniq in uniq:
                if item_data[1] >= item_uniq:
                    result[item_uniq] = result[item_uniq] + item_data[0]

        del result[0]  # удалить суммарную вероятность где ущерб 0
        return result

    def setup_page_format(self):
        """Настройка формата страницы А3 горизонтальный"""
        section = self.doc.sections[0]

        # Размер А3 в дюймах (297x420 мм)
        # 1 дюйм = 914400 твипов
        section.page_height = int(11.69 * 914400)  # 297мм
        section.page_width = int(16.54 * 914400)  # 420мм

        # Устанавливаем горизонтальную ориентацию
        section.orientation = WD_ORIENT.LANDSCAPE

        # Поля страницы (1 дюйм = 914400 твипов)
        margin = int(0.8 * 914400)  # 20мм
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin

    def generate_full_report(self, file_path: str, project_code: str = None):
        """Генерация полного отчета"""
        # Настраиваем формат страницы
        self.setup_page_format()

        # Заголовок документа
        self.doc.add_heading('Отчет по результатам расчетов', 0)

        # Таблица результатов расчета
        self.add_calculation_results_table(project_code)

        # Явно добавляем разрыв страницы перед анализом риска
        self.doc.add_page_break()

        # Добавляем раздел анализа риска
        self.add_risk_analysis(project_code)

        # Сохраняем документ
        self.doc.save(file_path)

    def add_risk_analysis(self, project_code: str = None):
        """Добавление раздела анализа риска"""
        self.doc.add_heading('Анализ риска', level=1)

        # Добавляем статистику
        self.add_risk_statistics(project_code)

        # Явно добавляем разрыв страницы перед диаграммами
        self.doc.add_page_break()

        # Добавляем F/N и F/G диаграммы
        self.add_fn_fg_diagrams(project_code)

    def add_calculation_results_table(self, project_code: str = None):
        """Добавление таблицы результатов расчета"""
        self.doc.add_heading('Результаты расчета', level=1)

        # Получаем результаты
        if project_code:
            results = self.calc_repo.get_by_project(project_code)
        else:
            results = self.calc_repo.get_all()

        if not results:
            self.doc.add_paragraph('Нет данных для отображения')
            return
        """Добавление подробной таблицы результатов"""

        # Получаем все результаты
        results = self.calc_repo.get_by_project(project_code) if project_code else self.calc_repo.get_all()
        # Сортируем результаты по номеру сценария
        results = sorted(results, key=lambda x: int(x.scenario_number))

        if not results:
            self.doc.add_paragraph('Нет данных для отображения')
            return

        # Создаем таблицу
        headers = [
            "№ п/п", "Код проекта", "№ сценария", "Оборудование",
            "Тип оборудования", "Тип вещества",
            "q_10.5 (кВт/м2)", "q_7.0 (кВт/м2)", "q_4.2 (кВт/м2)", "q_1.4 (кВт/м2)",
            "p_53 (кПа)", "p_28 (кПа)", "p_12 (кПа)", "p_5 (кПа)", "p_2 (кПа)",
            "Длина факела (м)", "Диаметр факела (м)",
            "Радиус НКПР (м)", "Радиус вспышки (м)",
            "Погибшие (чел)", "Пострадавшие (чел)",
            "Прямые потери (млн.руб)", "Затраты на ЛЛА (млн.руб)",
            "Социальный ущерб (млн.руб)", "Косвенный ущерб (млн.руб)",
            "Экологический ущерб (млн.руб)", "Суммарный ущерб (млн.руб)",
            "Риск гибели (чел/год)", "Риск травмирования (чел/год)",
            "Ожидаемый ущерб (млн.руб/год)", "Вероятность (1/год)",
            "Масса в аварии (т)", "Масса в ПФ (т)", "Масса в оборудовании (т)"
        ]

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Заполняем заголовки
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


        # Заполняем данные таблицы
        for i, result in enumerate(results, 1):
            row_cells = table.add_row().cells

            # Базовая информация
            row_cells[0].text = str(i)
            row_cells[1].text = result.project_code
            row_cells[2].text = str(result.scenario_number)
            row_cells[3].text = result.equipment_name
            row_cells[4].text = EquipmentType.get_display_name(result.equipment_type)
            row_cells[5].text = SubstanceType.get_display_name(result.substance_type)

            # Параметры теплового излучения
            row_cells[6].text = f"{result.q_10_5:.2f}" if result.q_10_5 else "-"
            row_cells[7].text = f"{result.q_7_0:.2f}" if result.q_7_0 else "-"
            row_cells[8].text = f"{result.q_4_2:.2f}" if result.q_4_2 else "-"
            row_cells[9].text = f"{result.q_1_4:.2f}" if result.q_1_4 else "-"

            # Параметры ударной волны
            row_cells[10].text = f"{result.p_53:.2f}" if result.p_53 else "-"
            row_cells[11].text = f"{result.p_28:.2f}" if result.p_28 else "-"
            row_cells[12].text = f"{result.p_12:.2f}" if result.p_12 else "-"
            row_cells[13].text = f"{result.p_5:.2f}" if result.p_5 else "-"
            row_cells[14].text = f"{result.p_2:.2f}" if result.p_2 else "-"

            # Параметры факела
            row_cells[15].text = f"{result.l_f:.2f}" if result.l_f else "-"
            row_cells[16].text = f"{result.d_f:.2f}" if result.d_f else "-"

            # Параметры вспышки
            row_cells[17].text = f"{result.r_nkpr:.2f}" if result.r_nkpr else "-"
            row_cells[18].text = f"{result.r_flash:.2f}" if result.r_flash else "-"

            # Последствия
            row_cells[19].text = str(result.casualties)
            row_cells[20].text = str(result.injured)

            # Ущерб
            row_cells[21].text = f"{result.direct_losses:.2f}"
            row_cells[22].text = f"{result.liquidation_costs:.2f}"
            row_cells[23].text = f"{result.social_losses:.2f}"
            row_cells[24].text = f"{result.indirect_damage:.2f}"
            row_cells[25].text = f"{result.environmental_damage:.2f}"
            row_cells[26].text = f"{result.total_damage:.2f}"

            # Риски
            row_cells[27].text = f"{result.casualty_risk:.2e}"
            row_cells[28].text = f"{result.injury_risk:.2e}"
            row_cells[29].text = f"{result.expected_damage:.2e}"
            row_cells[30].text = f"{result.probability:.2e}"

            # Массы
            row_cells[31].text = f"{result.mass_in_accident:.2f}"
            row_cells[32].text = f"{result.mass_in_factor:.2f}"
            row_cells[33].text = f"{result.mass_in_equipment:.2f}"

    # report_generator.py (продолжение)

    def add_risk_statistics(self, project_code: str = None):
        """Добавление статистических показателей"""
        results = self.calc_repo.get_by_project(project_code) if project_code else self.calc_repo.get_all()

        self.doc.add_heading('Статистические показатели', level=2)

        stats_table = self.doc.add_table(rows=8, cols=2)
        stats_table.style = 'Table Grid'

        # Заполняем статистику
        rows = stats_table.rows
        rows[0].cells[0].text = "Всего сценариев"
        rows[0].cells[1].text = str(len(results))

        rows[1].cells[0].text = "Максимальное число погибших"
        rows[1].cells[1].text = str(max(r.casualties for r in results))

        rows[2].cells[0].text = "Максимальное число пострадавших"
        rows[2].cells[1].text = str(max(r.injured for r in results))

        rows[3].cells[0].text = "Максимальный ущерб (млн.руб)"
        rows[3].cells[1].text = f"{max(r.total_damage for r in results):.2f}"

        rows[4].cells[0].text = "Суммарный риск гибели (чел/год)"
        rows[4].cells[1].text = f"{sum(r.casualty_risk for r in results):.2e}"

        rows[5].cells[0].text = "Суммарный риск травмирования (чел/год)"
        rows[5].cells[1].text = f"{sum(r.injury_risk for r in results):.2e}"

        rows[6].cells[0].text = "Максимальная частота аварий с гибелью (1/год)"
        death_scenarios = [r for r in results if r.casualties >= 1]
        max_death_frequency = max((r.probability for r in death_scenarios), default=0)
        rows[6].cells[1].text = f"{max_death_frequency:.2e}"

        rows[7].cells[0].text = "Максимальный экологический ущерб (млн.руб)"
        rows[7].cells[1].text = f"{max(r.environmental_damage for r in results):.2f}"

    # report_generator.py (продолжение)

    def add_fn_fg_diagrams(self, project_code: str = None):
        """Добавление F/N и F/G диаграмм"""
        self.doc.add_heading('F/N и F/G диаграммы', level=2)

        results = self.calc_repo.get_by_project(project_code) if project_code else self.calc_repo.get_all()

        # Создаем F/N диаграмму
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 5))

        # F/N диаграмма
        # Построение Fn диаграммы
        casualty_data = []
        for result in results:
            if result.casualties > 0:
                casualty_data.append((result.probability, result.casualties))

        sum_data = self._sum_data_for_fn(casualty_data)

        if casualty_data:
            people, probability = list(sum_data.keys()), list(sum_data.values())
            # для сплошных горизонтальных линий
            chart_line_x = []
            chart_line_y = []

            for i in range(len(people) - 1):
                # Добавляем две точки для текущего горизонтального отрезка
                chart_line_x.extend([people[i], people[i + 1], None])
                chart_line_y.extend([probability[i], probability[i], None])

            print(chart_line_x, chart_line_y)
            # для вертикальных пунктирных линий
            chart_dot_line_x = []
            chart_dot_line_y = []

            for i in range(len(people) - 1):
                # Добавляем две точки для вертикального пунктирного отрезка
                chart_dot_line_x.extend([people[i + 1], people[i + 1]])
                chart_dot_line_y.extend([probability[i], probability[i + 1]])

            # Строим F/N кривую
            ax1.semilogy(chart_line_x, chart_line_y, color='b', linestyle='-', marker='.')
            ax1.semilogy(chart_dot_line_x, chart_dot_line_y, color='b', linestyle='--', marker='.')
            ax1.set_xticks(people)
            ax1.grid(True)
            ax1.set_xlabel('N, число погибших')
            ax1.set_ylabel('F, частота событий с N и более погибшими, 1/год')
            ax1.set_title('F/N диаграмма')

        # F/G диаграмма
        damage_data = []
        for result in results:
            if result.casualties > 0:
                damage_data.append((result.probability, result.casualties))

        sum_data = self._sum_data_for_fg(damage_data)

        if damage_data:
            damage, probability = list(sum_data.keys()), list(sum_data.values())
            # для сплошных линий
            chart_line_x = []
            chart_line_y = []
            for i in damage:
                if damage[0] == i:
                    chart_line_x.extend([0, i, i, i])
                    chart_line_y.extend([probability[damage.index(i)], probability[damage.index(i)], None, None])
                elif damage[-1] == i:
                    chart_line_x.extend([damage[damage.index(i)-1], damage[damage.index(i)-1], i, i])
                    chart_line_y.extend([probability[damage.index(i)], probability[damage.index(i)], probability[damage.index(i)], probability[damage.index(i)]])
                    break
                else:
                    chart_line_x.extend([damage[damage.index(i) - 1], i, i, i])
                    chart_line_y.extend([probability[damage.index(i)], probability[damage.index(i)], None, None])

            # для пунктирных линий
            chart_dot_line_x = []
            chart_dot_line_y = []
            for i in damage:
                if i == damage[-1]:
                    chart_dot_line_x.extend([i, i])
                    chart_dot_line_y.extend([probability[damage.index(i)], probability[damage.index(i)]])
                    chart_dot_line_x.extend([i, i])
                    chart_dot_line_y.extend([probability[damage.index(i)], 0])
                    break
                chart_dot_line_x.extend([i, i])
                chart_dot_line_y.extend([probability[damage.index(i)], probability[damage.index(i) + 1]])

            # Создание графика
            # Построение основной диаграммы
            ax2.semilogy(chart_line_x, chart_line_y, color='r', linestyle='-', marker='.')
            ax2.semilogy(chart_dot_line_x, chart_dot_line_y, color='r', linestyle='--', marker='.')
            ax2.grid(True)
            ax2.set_xticks(damage)
            ax2.set_xlabel('G, ущерб, млн.руб')
            ax2.set_ylabel('F, частота событий с ущербом G и более, 1/год')
            ax2.set_title('F/G диаграмма')

        plt.tight_layout()

        # Сохраняем диаграммы во временный буфер
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)

        # Добавляем изображение в документ
        self.doc.add_picture(img_buffer, width=Inches(7))

        # Закрываем графики
        plt.close()