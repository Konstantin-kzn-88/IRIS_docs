from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def find_paragraph_with_marker(doc: Document, marker: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if marker in p.text:
            return p
    return None


def clear_paragraph(p: Paragraph) -> None:
    for r in p.runs[::-1]:
        p._p.remove(r._r)
    p.text = ""


def insert_paragraph_after(doc: Document, p: Paragraph, text: str = "") -> Paragraph:
    """
    Совместимо со старыми python-docx: создаём абзац и переносим XML сразу после p.
    """
    new_p = doc.add_paragraph(text)
    p._p.addnext(new_p._p)
    return new_p


def insert_table_after(doc: Document, p: Paragraph, rows: int, cols: int, style: str = "Table Grid"):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    p._p.addnext(table._tbl)
    if rows and rows > 0:
        _set_repeat_table_header(table, header_rows=1)
    return table




def _set_repeat_table_header(table, header_rows: int = 1) -> None:
    """
    Делает первые header_rows строк таблицы повторяемыми (как заголовок на каждой странице).
    Работает через OOXML, совместимо со старыми версиями python-docx.
    """
    try:
        for i in range(min(header_rows, len(table.rows))):
            tr = table.rows[i]._tr
            trPr = tr.get_or_add_trPr()
            # удалить существующие tblHeader, если были
            for el in trPr.findall(qn('w:tblHeader')):
                trPr.remove(el)
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), 'true')
            trPr.append(tblHeader)
    except Exception:
        # не падать, если какая-то версия docx не поддерживает элементы
        return


def insert_paragraph_after_table(doc: Document, table, text: str = "") -> Paragraph:
    p = doc.add_paragraph(text)
    table._tbl.addnext(p._p)
    return p


def add_section_header_row(table, title: str):
    row = table.add_row().cells
    row[0].merge(row[1])

    p = row[0].paragraphs[0]
    run = p.add_run(title)
    set_run_font(run, bold=True)

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_run_font(run, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    run.bold = bold