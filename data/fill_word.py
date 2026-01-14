import json
import math
from pathlib import Path
from docx.shared import Cm
from docx import Document

from report.constants import NGK_BACKGROUND_RISK
from report.formatters import risk_to_dbr, risk_to_ppm
from report.constants import SOCIAL_FATALITY_RISKS_DBR
from report.paths import BASE_DIR, DB_PATH, TEMPLATE_PATH, OUT_PATH
from report.db import (
    open_db,
    get_used_substances,
    get_used_equipment,
    get_hazard_distribution,
    get_scenarios,
    get_ov_amounts_in_accident,
    get_impact_zones,
    get_personnel_casualties,
    get_damage,
    get_collective_risk,
    get_individual_risk,
    get_fatal_accident_frequency_range,
    get_max_damage_by_hazard_component,
    get_fn_source_rows,
    get_fg_source_rows,
    get_pareto_risk_source_rows,
    get_pareto_damage_source_rows,
    get_pareto_environmental_damage_source_rows,
    get_max_losses_by_hazard_component,
    get_risk_matrix_rows,
    get_risk_matrix_damage_rows,
    get_top_scenarios_by_hazard_component,
    get_damage_by_component,
    get_substances_by_component,
)
from report.sections import SUBSTANCE_SECTIONS, EQUIPMENT_SECTIONS
from report.formatters import (
    format_value,
    pretty_json_substance,
    pretty_json_generic,
    format_exp,
    format_float_3,
    format_float_1,
)
from report.word_utils import (
    find_paragraph_with_marker,
    clear_paragraph,
    insert_paragraph_after,
    insert_table_after,
    insert_paragraph_after_table,
    add_section_header_row,
    set_run_font,
)

from report.charts import (
    build_fn_points,
    build_fg_points,
    save_fn_chart,
    save_fg_chart,
    build_pareto_series,
    save_pareto_chart,
    save_component_damage_chart,
    save_risk_matrix_chart,
    save_risk_matrix_chart_damage,
)


def set_cell_text(cell, text, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    set_run_font(run, bold=bold)


def fill_table(table, obj: dict, sections, json_formatter):
    for section_title, fields in sections:
        rows = []
        for field, label in fields:
            raw = obj.get(field)

            if field.endswith("_json"):
                value = json_formatter(raw)
            else:
                value = format_value(field, raw)

            # Не выводим пустое. "-" не выводим (если нужно выводить — убери "-")
            if value is None:
                continue
            if isinstance(value, str) and value.strip() in ("", "-"):
                continue

            rows.append((label, value))

        if not rows:
            continue

        add_section_header_row(table, section_title)
        for label, value in rows:
            r = table.add_row().cells
            set_cell_text(r[0], label)
            set_cell_text(r[1], value)


def fmt(value):
    """Для таблицы зон поражающих факторов:
    - None -> '-'
    - если |value| >= 1: округлить до целого (int)
    - если |value| < 1: округлить до 1 знака после запятой
    """
    if value is None:
        return "-"

    try:
        v = float(value)
    except Exception:
        return str(value)

    if abs(v) >= 1:
        return str(int(round(v)))
    return f"{v:.1f}"


def render_section_at_marker(
        doc: Document,
        marker: str,
        section_title: str,
        items: list[dict],
        item_title_field: str,
        sections,
        json_formatter,
):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, section_title)
    if p.runs:
        set_run_font(p.runs[0], bold=True)

    cur = p
    for item in items:
        item_title = item.get(item_title_field, "—")
        cur = insert_paragraph_after(doc, cur, str(item_title))
        if cur.runs:
            set_run_font(cur.runs[0], bold=True)

        table = insert_table_after(doc, cur, rows=0, cols=2, style="Table Grid")
        fill_table(table, item, sections, json_formatter)

        cur = insert_paragraph_after_table(doc, table, "")


def render_distribution_table_at_marker(doc: Document, marker: str, title: str, rows: list[dict]):
    """
    Горизонтальная таблица:
    1) Наименование оборудования
    2) Наименование вещества
    3) Количество опасного вещества, т
    4) Агрегатное состояние
    5) Давление, МПа
    6) Температура, °C

    Требование: округление чисел до 3 знаков после запятой.
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, title)
    if p.runs:
        set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=6, style="Table Grid")
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Наименование оборудования", bold=True)
    set_cell_text(hdr[1], "Наименование вещества", bold=True)
    set_cell_text(hdr[2], "Количество опасного вещества, т", bold=True)
    set_cell_text(hdr[3], "Агрегатное состояние", bold=True)
    set_cell_text(hdr[4], "Давление, МПа", bold=True)
    set_cell_text(hdr[5], "Температура, °C", bold=True)

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("equipment_name") or "-")
        set_cell_text(row[1], r.get("substance_name") or "-")
        set_cell_text(row[2], format_float_3(r.get("amount_t")))
        set_cell_text(row[3], r.get("phase_state") if r.get("phase_state") is not None else "-")
        set_cell_text(row[4], format_float_3(r.get("pressure_mpa")))
        set_cell_text(row[5], format_float_3(r.get("substance_temperature_c")))

    insert_paragraph_after_table(doc, table, "")


def render_ov_amount_table_at_marker(doc: Document, marker: str, title: str, rows: list[dict]):
    """
    Таблица: Оценка количества опасного вещества в аварии
    Колонки:
      1) № п/п
      2) Наименование оборудования
      3) Номер сценария (С{scenario_no})
      4) Количество ОВ участвующего в аварии, т (ov_in_accident_t)
      5) Количество ОВ в создании поражающего фактора, т (ov_in_hazard_factor_t)

    Требование: округление чисел до 3 знаков после запятой.
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, title)
    if p.runs:
        set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=5, style="Table Grid")

    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "№ п/п", bold=True)
    set_cell_text(hdr[1], "Наименование оборудования", bold=True)
    set_cell_text(hdr[2], "Номер сценария", bold=True)
    set_cell_text(hdr[3], "Количество ОВ участвующего в аварии, т", bold=True)
    set_cell_text(hdr[4], "Количество ОВ в создании поражающего фактора, т", bold=True)

    # Порядок С1..Сn, внутри номера — по оборудованию
    rows_sorted = sorted(
        rows,
        key=lambda r: (
            int(r.get("scenario_no")) if r.get("scenario_no") is not None else 10 ** 9,
            str(r.get("equipment_name") or ""),
        ),
    )

    for idx, r in enumerate(rows_sorted, start=1):
        row = table.add_row().cells
        set_cell_text(row[0], idx)
        set_cell_text(row[1], r.get("equipment_name") or "-")
        sc_no = r.get("scenario_no")
        set_cell_text(row[2], f"С{sc_no}" if sc_no is not None else "-")
        set_cell_text(row[3], format_float_3(r.get("ov_in_accident_t")))
        set_cell_text(row[4], format_float_3(r.get("ov_in_hazard_factor_t")))

    insert_paragraph_after_table(doc, table, "")


def render_personnel_casualties_table_at_marker(doc: Document, marker: str, title: str, rows: list[dict]):
    """Таблица: Оценка количества погибших/пострадавших.

    Колонки из calculations:
      1) № п/п
      2) Наименование оборудования (equipment_name)
      3) Номер сценария (С{scenario_no})
      4) Количество погибших, чел (fatalities_count)
      5) Количество пострадавших, чел (injured_count)

    Требование: None заменить на "-".
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, title)
    if p.runs:
        set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=5, style="Table Grid")

    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "№ п/п", bold=True)
    set_cell_text(hdr[1], "Наименование оборудования", bold=True)
    set_cell_text(hdr[2], "Номер сценария", bold=True)
    set_cell_text(hdr[3], "Количество погибших, чел", bold=True)
    set_cell_text(hdr[4], "Количество пострадавших, чел", bold=True)

    # Порядок С1..Сn, внутри номера — по оборудованию
    rows_sorted = sorted(
        rows,
        key=lambda r: (
            int(r.get("scenario_no")) if r.get("scenario_no") is not None else 10 ** 9,
            str(r.get("equipment_name") or ""),
        ),
    )

    for idx, r in enumerate(rows_sorted, start=1):
        sc_no = r.get("scenario_no")
        row = table.add_row().cells
        set_cell_text(row[0], idx)
        set_cell_text(row[1], r.get("equipment_name") or "-")
        set_cell_text(row[2], f"С{sc_no}" if sc_no is not None else "-")

        fat = r.get("fatalities_count")
        inj = r.get("injured_count")
        set_cell_text(row[3], fat if fat is not None else "-")
        set_cell_text(row[4], inj if inj is not None else "-")

    insert_paragraph_after_table(doc, table, "")


def _load_typical_scenarios() -> dict:
    """typical_scenarios.json не изменяем; читаем из стандартных мест."""
    candidates = [
        BASE_DIR / "calc" / "typical_scenarios.json",
        BASE_DIR / "typical_scenarios.json",
    ]
    for p in candidates:
        if p.exists():
            with open(p, "r", encoding="utf-8") as f:
                return json.load(f)
    return {}


def _scenario_item_to_text(item) -> str:
    if isinstance(item, str):
        return item
    if isinstance(item, dict):
        # поддержка {"scenario_text": "..."} и близких вариантов
        return str(item.get("scenario_text") or item.get("text") or item.get("description") or "Описание не задано")
    return "Описание не задано"


def _get_scenarios_root(typical: dict) -> dict:
    # root может быть либо сразу словарём, либо лежать в ключе "scenarios"
    if isinstance(typical, dict) and "scenarios" in typical and isinstance(typical["scenarios"], dict):
        return typical["scenarios"]
    return typical if isinstance(typical, dict) else {}


def _get_description_list(root: dict, equipment_type, kind):
    # основной формат: root[equipment_type][kind] -> list
    et = str(equipment_type)
    kd = str(kind)
    if et in root and isinstance(root[et], dict) and kd in root[et]:
        return root[et][kd]
    # запасной вариант: root[kind][equipment_type]
    if kd in root and isinstance(root[kd], dict) and et in root[kd]:
        return root[kd][et]
    return None


def _build_scenario_index(rows: list[dict]):
    """
    Для каждой пары (equipment_type, kind) строим отображение:
    scenario_no (уникальный, отсортированный) -> локальный индекс 0..N-1
    """
    mapping = {}
    for r in rows:
        key = (r.get("equipment_type"), r.get("substance_kind"))
        mapping.setdefault(key, set()).add(r.get("scenario_no"))

    idx_map = {}
    for key, sc_set in mapping.items():
        sc_list = sorted([x for x in sc_set if x is not None])
        idx_map[key] = {sc_no: i for i, sc_no in enumerate(sc_list)}
    return idx_map


def render_scenarios_table_at_marker(doc: Document, marker: str, title: str, rows: list[dict]):
    """
    Таблица сценариев:
    1) № п/п
    2) Наименование оборудования
    3) Номер сценария (С{scenario_no})
    4) Описание сценария (typical_scenarios.json по (equipment_type, kind) и позиции)
    5) Базовая частота, 1/год (base_frequency, экспонента)
    6) Условная вероятность, -
    7) Частота сценария аварии, 1/год (scenario_frequency, экспонента)
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, title)
    if p.runs:
        set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=7, style="Table Grid")
    hdr = table.rows[0].cells

    headers = [
        "№ п/п",
        "Наименование оборудования",
        "Номер сценария",
        "Описание сценария",
        "Базовая частота, 1/год",
        "Условная вероятность, -",
        "Частота сценария аварии, 1/год",
    ]
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)

    typical = _load_typical_scenarios()
    root = _get_scenarios_root(typical)
    idx_map = _build_scenario_index(rows)

    # Порядок: С1..Сn, внутри — по оборудованию
    rows_sorted = sorted(
        rows,
        key=lambda r: (
            int(r.get("scenario_no")) if r.get("scenario_no") is not None else 10 ** 9,
            str(r.get("equipment_name") or ""),
        ),
    )

    for i, r in enumerate(rows_sorted, start=1):
        et = r.get("equipment_type")
        kd = r.get("substance_kind")
        sc_no = r.get("scenario_no")

        local_idx = idx_map.get((et, kd), {}).get(sc_no, None)
        desc_list = _get_description_list(root, et, kd)
        if isinstance(desc_list, list) and local_idx is not None and 0 <= local_idx < len(desc_list):
            desc = _scenario_item_to_text(desc_list[local_idx])
        else:
            desc = "Описание не задано"

        row = table.add_row().cells
        set_cell_text(row[0], i)
        set_cell_text(row[1], r.get("equipment_name") or "-")
        set_cell_text(row[2], f"С{sc_no}" if sc_no is not None else "-")
        set_cell_text(row[3], desc)
        set_cell_text(row[4], format_exp(r.get("base_frequency")))
        set_cell_text(row[5],
                      r.get("accident_event_probability") if r.get("accident_event_probability") is not None else "-")
        set_cell_text(row[6], format_exp(r.get("scenario_frequency")))

    insert_paragraph_after_table(doc, table, "")


def render_impact_zones_table(doc, marker: str, rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, "Зоны действия поражающих факторов")
    set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=23, style="Table Grid")

    headers = [
        "№ п/п",
        "Наименование оборудования",
        "Номер сценария",

        "q = 10,5",
        "q = 7,0",
        "q = 4,2",
        "q = 1,4",

        "ΔР = 70",
        "ΔР = 28",
        "ΔР = 14",
        "ΔР = 5",
        "ΔР = 2",

        "Lf",
        "Df",
        "Rнкпр",
        "Rвсп",
        "Rlpt",
        "Rppt",

        "Q600",
        "Q320",
        "Q220",
        "Q120",

        "St",
    ]

    # шапка
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    # повторяемая шапка (через word_utils, у тебя уже работает)
    # ничего дописывать не нужно

    for idx, r in enumerate(rows, start=1):
        row = table.add_row().cells

        set_cell_text(row[0], idx)
        set_cell_text(row[1], r["equipment_name"])
        set_cell_text(row[2], f"С{r['scenario_no']}")

        set_cell_text(row[3], fmt(r["q_10_5"]))
        set_cell_text(row[4], fmt(r["q_7_0"]))
        set_cell_text(row[5], fmt(r["q_4_2"]))
        set_cell_text(row[6], fmt(r["q_1_4"]))

        set_cell_text(row[7], fmt(r["p_70"]))
        set_cell_text(row[8], fmt(r["p_28"]))
        set_cell_text(row[9], fmt(r["p_14"]))
        set_cell_text(row[10], fmt(r["p_5"]))
        set_cell_text(row[11], fmt(r["p_2"]))

        set_cell_text(row[12], fmt(r["l_f"]))
        set_cell_text(row[13], fmt(r["d_f"]))
        set_cell_text(row[14], fmt(r["r_nkpr"]))
        set_cell_text(row[15], fmt(r["r_vsp"]))
        set_cell_text(row[16], fmt(r["l_pt"]))
        set_cell_text(row[17], fmt(r["p_pt"]))

        set_cell_text(row[18], fmt(r["q_600"]))
        set_cell_text(row[19], fmt(r["q_320"]))
        set_cell_text(row[20], fmt(r["q_220"]))
        set_cell_text(row[21], fmt(r["q_120"]))

        set_cell_text(row[22], fmt(r["s_t"]))


def render_damage_table_at_marker(doc, marker: str, rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, "Ущерб")
    set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=9, style="Table Grid")

    headers = [
        "№ п/п",
        "Наименование оборудования",
        "Номер сценария",
        "Прямые потери, тыс.руб",
        "Затраты на ликвидацию, тыс.руб",
        "Социальные потери, тыс.руб",
        "Косвенный ущерб, тыс.руб",
        "Экологический ущерб, тыс.руб",
        "Суммарный ущерб, тыс.руб",
    ]

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for idx, r in enumerate(rows, start=1):
        row = table.add_row().cells
        set_cell_text(row[0], idx)
        set_cell_text(row[1], r.get("equipment_name", "-"))
        set_cell_text(row[2], f"С{r.get('scenario_no')}" if r.get("scenario_no") is not None else "-")

        set_cell_text(row[3], format_float_1(r.get("direct_losses")))
        set_cell_text(row[4], format_float_1(r.get("liquidation_costs")))
        set_cell_text(row[5], format_float_1(r.get("social_losses")))
        set_cell_text(row[6], format_float_1(r.get("indirect_damage")))
        set_cell_text(row[7], format_float_1(r.get("total_environmental_damage")))
        set_cell_text(row[8], format_float_1(r.get("total_damage")))

    insert_paragraph_after_table(doc, table, "")


def render_collective_risk_table(doc, marker: str, rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, "Коллективный риск")
    set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=3, style="Table Grid")

    headers = [
        "Составляющая ОПО",
        "Коллективный риск гибели, чел·год⁻¹",
        "Коллективный риск ранения, чел·год⁻¹",
    ]

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("hazard_component", "-"))
        set_cell_text(row[1], format_exp(r.get("collective_risk_fatalities")))
        set_cell_text(row[2], format_exp(r.get("collective_risk_injured")))

    insert_paragraph_after_table(doc, table, "")


def render_individual_risk_table(doc, marker: str, rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, "Индивидуальный риск")
    set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=3, style="Table Grid")

    headers = [
        "Составляющая ОПО",
        "Индивидуальный риск гибели, 1·год⁻¹",
        "Индивидуальный риск ранения, 1·год⁻¹",
    ]

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("hazard_component", "-"))
        set_cell_text(row[1], format_exp(r.get("individual_risk_fatalities")))
        set_cell_text(row[2], format_exp(r.get("individual_risk_injured")))

    insert_paragraph_after_table(doc, table, "")


def render_fatal_accident_frequency_text(doc, marker: str, min_freq, max_freq):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    if min_freq is None or max_freq is None:
        text = "Частота аварии с гибелью не менее одного человека равна 0."
        p = insert_paragraph_after(doc, p_marker, text)
        set_run_font(p.runs[0], bold=False)
        return

    text = (
        "Частота аварии с гибелью не менее одного человека "
        f"лежит в диапазоне для объекта от {format_exp(min_freq)} "
        f"и до {format_exp(max_freq)} 1/год."
    )

    p = insert_paragraph_after(doc, p_marker, text)
    set_run_font(p.runs[0], bold=False)


def render_max_damage_by_component_table(doc, marker: str, rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, "Ущерб имуществу и окружающей среде")
    set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=3, style="Table Grid")

    headers = [
        "Составляющая ОПО",
        "Максимальный суммарный ущерб, тыс.руб",
        "Максимальный экологический ущерб, тыс.руб",
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("hazard_component", "-"))
        set_cell_text(row[1], format_float_1(r.get("max_total_damage")))
        set_cell_text(row[2], format_float_1(r.get("max_total_environmental_damage")))

    insert_paragraph_after_table(doc, table, "")


def render_chart_at_marker(doc: Document, marker: str, title: str, image_path: Path | None, width_cm: float = 16.0):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, title)
    if p.runs:
        set_run_font(p.runs[0], bold=True)

    if image_path is None or not Path(image_path).exists():
        p2 = insert_paragraph_after(doc, p, "Данные для построения диаграммы отсутствуют.")
        if p2.runs:
            set_run_font(p2.runs[0], bold=False)
        return

    pic_p = insert_paragraph_after(doc, p, "")
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def render_fn_chart_at_marker(doc: Document, marker: str, fn_rows: list[dict]):
    points = build_fn_points(fn_rows)
    if not points:
        image_path = None
    else:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "fn.png"
        save_fn_chart(points, image_path)

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="F/N - диаграмма",
        image_path=image_path,
        width_cm=16.0,
    )


def render_fg_chart_at_marker(doc: Document, marker: str, fg_rows: list[dict]):
    points = build_fg_points(fg_rows)
    if not points:
        image_path = None
    else:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "fg.png"
        save_fg_chart(points, image_path)

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="F/G - диаграмма",
        image_path=image_path,
        width_cm=16.0,
    )


def render_pareto_fatalities_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    """Pareto по коллективному риску гибели (collective_risk_fatalities)"""
    series = build_pareto_series(rows, "collective_risk_fatalities")
    charts_dir = OUT_PATH.parent / "charts"
    charts_dir.mkdir(parents=True, exist_ok=True)
    img_path = charts_dir / "pareto_fatalities.png" if series else None
    if series:
        save_pareto_chart(
            series,
            img_path,
            title="Pareto-диаграмма (вклад) сценариев по коллективному риску гибели",
            ylabel="Коллективный риск гибели, чел·год⁻¹",
        )
    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Pareto сценариев по коллективному риску гибели",
        image_path=img_path,
        width_cm=16.0,
    )


def render_pareto_injured_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    """Pareto по коллективному риску ранения (collective_risk_injured)"""
    series = build_pareto_series(rows, "collective_risk_injured")
    charts_dir = OUT_PATH.parent / "charts"
    charts_dir.mkdir(parents=True, exist_ok=True)
    img_path = charts_dir / "pareto_injured.png" if series else None
    if series:
        save_pareto_chart(
            series,
            img_path,
            title="Pareto-диаграмма (вклад) сценариев по коллективному риску ранения",
            ylabel="Коллективный риск ранения, чел·год⁻¹",
        )
    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Pareto сценариев по коллективному риску ранения",
        image_path=img_path,
        width_cm=16.0,
    )


def render_pareto_damage_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    # собираем серии: подпись = "equipment / Сn", значение = total_damage
    series = build_pareto_series(rows, value_key="total_damage")
    image_path = None

    if series:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "pareto_damage.png"

        save_pareto_chart(
            series=series,
            path=image_path,
            title="Pareto-диаграмма (вклад) сценариев по суммарному ущербу",
            ylabel="Суммарный ущерб, тыс.руб",
        )

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Pareto сценариев по суммарному ущербу",
        image_path=image_path,
        width_cm=16.0,
    )


def render_pareto_environmental_damage_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    series = build_pareto_series(rows, value_key="total_environmental_damage")
    image_path = None

    if series:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "pareto_environmental_damage.png"

        save_pareto_chart(
            series=series,
            path=image_path,
            title="Pareto-диаграмма (вклад) сценариев по экологическому ущербу",
            ylabel="Экологический ущерб, тыс.руб",
        )

    render_chart_at_marker(
        doc,
        marker,
        "Pareto сценариев по экологическому ущербу",
        image_path,
    )


def render_component_damage_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    image_path = None
    if rows:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "damage_by_component.png"
        save_component_damage_chart(rows, image_path)

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Распределение ущерба по составляющим ОПО",
        image_path=image_path,
        width_cm=16.0,
    )


def render_risk_matrix_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    image_path = None
    if rows:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "risk_matrix.png"
        save_risk_matrix_chart(rows, image_path)

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Матрица риска (частота – последствия)",
        image_path=image_path,
        width_cm=16.0,
    )


def render_risk_matrix_damage_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    image_path = None
    if rows:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "risk_matrix_damage.png"
        save_risk_matrix_chart_damage(rows, image_path)

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Матрица риска (частота – ущерб)",
        image_path=image_path,
        width_cm=16.0,
    )


def render_top_scenarios_by_component_table(doc, marker: str, rows: list[dict]):
    """
    Колонки:
      1) Составляющая объекта (hazard_component)
      2) Тип сценария (наиболее опасный / наиболее вероятный)
      3) Номер сценария (С{scenario_no})
      4) Оборудование (equipment_name)
      5) Количество погибших, чел (fatalities_count)
      6) Ущерб, тыс.руб (total_damage)
      7) Частота сценария, 1/год (scenario_frequency)
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker,
                               "Наиболее опасные и наиболее вероятные сценарии аварии по составляющим объекта")
    set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=7, style="Table Grid")

    headers = [
        "Составляющая объекта",
        "Тип сценария",
        "Номер сценария",
        "Оборудование",
        "Количество погибших, чел",
        "Ущерб, тыс.руб",
        "Частота сценария, 1/год",
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    type_map = {
        "dangerous": "Наиболее опасный",
        "probable": "Наиболее вероятный",
    }

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("hazard_component", "-"))
        set_cell_text(row[1], type_map.get(r.get("scenario_type"), str(r.get("scenario_type", "-"))))

        sc_no = r.get("scenario_no")
        set_cell_text(row[2], f"С{sc_no}" if sc_no is not None else "-")

        set_cell_text(row[3], r.get("equipment_name", "-"))
        set_cell_text(row[4], r.get("fatalities_count", 0))
        set_cell_text(row[5], format_float_1(r.get("total_damage")))
        set_cell_text(row[6], format_exp(r.get("scenario_frequency")))

    insert_paragraph_after_table(doc, table, "")


def render_fatality_risk_by_component_table(doc, marker: str, rows: list[dict]):
    """Сводная таблица: индивидуальный и коллективный риск гибели по составляющим ОПО."""
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, "Коллективный и индивидуальный риск гибели по составляющим ОПО")
    set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=3, style="Table Grid")

    headers = [
        "Составляющая ОПО",
        "Индивидуальный риск гибели, 1·год⁻¹",
        "Коллективный риск гибели, чел·год⁻¹",
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("hazard_component", "-"))
        set_cell_text(row[1], format_exp(r.get("individual_risk_fatalities")))
        set_cell_text(row[2], format_exp(r.get("collective_risk_fatalities")))

    insert_paragraph_after_table(doc, table, "")


def render_comparative_fatality_risk_table(doc, marker: str, individual_risk_rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    title_p = insert_paragraph_after(
        doc, p_marker,
        "Сравнение риска гибели от различных причин и риска гибели при авариях на ОПО"
    )
    if title_p.runs:
        set_run_font(title_p.runs[0], bold=True)

    table = insert_table_after(doc, title_p, rows=1, cols=2, style="Table Grid")

    # Заголовки
    hdr = table.rows[0].cells
    hdr[0].text = ""
    r0 = hdr[0].paragraphs[0].add_run("Вид смертельной опасности")
    set_run_font(r0, bold=True)

    hdr[1].text = ""
    r1 = hdr[1].paragraphs[0].add_run("Уровень риска, дБR")
    set_run_font(r1, bold=True)

    # 1) Социальные риски
    for name, dbr in SOCIAL_FATALITY_RISKS_DBR:
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(str(name))
        row[1].paragraphs[0].add_run(f"{dbr:+.1f}")

    # 2) ОПО по составляющим (берем уже посчитанный индивидуальный риск)
    for r in individual_risk_rows:
        comp = r.get("hazard_component")
        # поддержка обоих возможных ключей
        risk = (
            r.get("individual_risk_fatalities")
            if r.get("individual_risk_fatalities") is not None
            else r.get("individual_risk")
        )

        dbr = risk_to_dbr(risk)
        dbr_txt = "—" if dbr is None else f"{dbr:+.1f}"

        row = table.add_row().cells
        row[0].paragraphs[0].add_run(f"Риск гибели при аварии на ОПО ({comp})")
        row[1].paragraphs[0].add_run(dbr_txt)


def render_ngk_background_comparison_table(doc, marker: str, conn):
    p = find_paragraph_with_marker(doc, marker)
    if p is None:
        return

    clear_paragraph(p)

    title = insert_paragraph_after(
        doc,
        p,
        "Сравнение фоновых показателей риска нефтегазового комплекса и показателей ОПО"
    )
    if title.runs:
        set_run_font(title.runs[0], bold=True)

    table = insert_table_after(doc, title, rows=1, cols=2, style="Table Grid")

    # Заголовки
    hdr = table.rows[0].cells
    h0 = hdr[0].paragraphs[0].add_run("Параметр")
    h1 = hdr[1].paragraphs[0].add_run("Значение")
    set_run_font(h0, bold=True)
    set_run_font(h1, bold=True)

    # --- A. Фон НГК ---
    for name, value in NGK_BACKGROUND_RISK:
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(name)
        row[1].paragraphs[0].add_run(value)

    # --- B. ОПО по составляющим ---
    risks = get_individual_risk(conn)
    damage_by_comp = get_damage_by_component(conn)

    for r in risks:
        comp = r["hazard_component"]
        risk = r["individual_risk_fatalities"]

        if risk is None or risk <= 0:
            continue

        dbr = risk_to_dbr(risk)
        ppm = risk * 1e6
        damage = damage_by_comp.get(comp)

        # Ущерб
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(
            f"Ущерб при аварии на ОПО ({comp}), млн руб"
        )
        row[1].paragraphs[0].add_run(
            f"{damage:.1f}" if damage is not None else "—"
        )

        # Риск дБR
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(
            f"Риск гибели при аварии на ОПО ({comp}), дБR"
        )
        row[1].paragraphs[0].add_run(f"{dbr:+.1f}")

        # Риск ppm
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(
            f"Риск гибели при аварии на ОПО ({comp}), ppm"
        )
        row[1].paragraphs[0].add_run(f"{ppm:.2f}")


def render_substances_by_component_table(doc, marker: str, conn):
    from report.word_utils import (
        find_paragraph_with_marker,
        clear_paragraph,
        insert_paragraph_after,
        insert_table_after,
        set_run_font,
    )

    p = find_paragraph_with_marker(doc, marker)
    if p is None:
        return

    clear_paragraph(p)

    title = insert_paragraph_after(
        doc,
        p,
        "Количество веществ по составляющим объекта"
    )
    if title.runs:
        set_run_font(title.runs[0], bold=True)

    table = insert_table_after(doc, title, rows=1, cols=2, style="Table Grid")

    # Заголовки
    hdr = table.rows[0].cells
    r0 = hdr[0].paragraphs[0].add_run("Составляющая объекта")
    r1 = hdr[1].paragraphs[0].add_run("Количество вещества, т")
    set_run_font(r0, bold=True)
    set_run_font(r1, bold=True)

    data = get_substances_by_component(conn)

    for comp, items in data.items():
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(str(comp))

        p_cell = row[1].paragraphs[0]
        p_cell.text = ""

        for i, (name, mass) in enumerate(items):
            if i > 0:
                p_cell.add_run("\n")
            p_cell.add_run(f"{name} — {mass:.3f} т")


def render_top_scenarios_description_by_component_table(doc: Document, marker: str, conn):
    """
    Таблица:
    - Составляющая объекта
    - Тип сценария (наиболее опасный / наиболее вероятный)
    - Номер сценария (Сn)
    - Наименование оборудования
    - Краткое описание сценария
    - Частота, 1/год
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    title_p = insert_paragraph_after(
        doc,
        p_marker,
        "Описание наиболее опасного и наиболее вероятного сценария аварии по составляющим объекта"
    )
    if title_p.runs:
        set_run_font(title_p.runs[0], bold=True)

    table = insert_table_after(doc, title_p, rows=1, cols=6, style="Table Grid")

    headers = [
        "Составляющая объекта",
        "Тип сценария",
        "Номер сценария",
        "Наименование оборудования",
        "Краткое описание сценария",
        "Частота, 1/год",
    ]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        clear_paragraph(cell.paragraphs[0])
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True)

    # 1) Берем топ-сценарии по составляющим
    top_rows = get_top_scenarios_by_hazard_component(conn)

    # 2) Берем "строки сценариев" для построения описания как в render_scenarios_table_at_marker
    scenario_rows = get_scenarios(conn)

    # 3) Строим индекс scenario_no -> local_idx по (equipment_type, kind)
    idx_map = _build_scenario_index(scenario_rows)

    # 4) Загружаем типовые сценарии
    typical = _load_typical_scenarios()
    root = _get_scenarios_root(typical)

    # 5) Мапа scenario_no -> row из get_scenarios (для equipment_type/kind)
    sc_map = {}
    for r in scenario_rows:
        sc_no = r.get("scenario_no")
        if sc_no is not None and sc_no not in sc_map:
            sc_map[sc_no] = r

    def _scenario_type_label(st: str) -> str:
        if st == "dangerous":
            return "Наиболее опасный"
        if st == "probable":
            return "Наиболее вероятный"
        return st or "-"

    for r in top_rows:
        comp = r.get("hazard_component")
        st = r.get("scenario_type")
        sc_no = r.get("scenario_no")
        eq_name = r.get("equipment_name")
        freq = r.get("scenario_frequency")

        # описание
        desc = "Описание не задано"
        base = sc_map.get(sc_no)
        if base:
            et = base.get("equipment_type")
            kd = base.get("substance_kind")
            local_idx = None
            if (et, kd) in idx_map and sc_no in idx_map[(et, kd)]:
                local_idx = idx_map[(et, kd)][sc_no]

            desc_list = _get_description_list(root, et, kd)
            if isinstance(desc_list, list) and local_idx is not None and 0 <= local_idx < len(desc_list):
                desc = _scenario_item_to_text(desc_list[local_idx])

        row = table.add_row().cells
        set_cell_text(row[0], comp if comp is not None else "-")
        set_cell_text(row[1], _scenario_type_label(st))
        set_cell_text(row[2], f"С{sc_no}" if sc_no is not None else "-")
        set_cell_text(row[3], eq_name or "-")
        set_cell_text(row[4], desc)
        set_cell_text(row[5], format_exp(freq))

    insert_paragraph_after_table(doc, table, "")






def main():
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    with open_db(DB_PATH) as conn:
        substances = get_used_substances(conn)
        equipment = get_used_equipment(conn)
        distribution = get_hazard_distribution(conn)
        scenarios = get_scenarios(conn)
        ov_amounts = get_ov_amounts_in_accident(conn)
        impact_zones = get_impact_zones(conn)
        casualties = get_personnel_casualties(conn)
        damage_rows = get_damage(conn)
        collective_risk_rows = get_collective_risk(conn)
        individual_risk_rows = get_individual_risk(conn)
        min_f, max_f = get_fatal_accident_frequency_range(conn)
        max_damage_rows = get_max_damage_by_hazard_component(conn)
        fn_rows = get_fn_source_rows(conn)
        fg_rows = get_fg_source_rows(conn)
        pareto_rows = get_pareto_risk_source_rows(conn)
        pareto_damage_rows = get_pareto_damage_source_rows(conn)
        pareto_env_rows = get_pareto_environmental_damage_source_rows(conn)
        component_damage_rows = get_max_losses_by_hazard_component(conn)
        risk_matrix_rows = get_risk_matrix_rows(conn)
        risk_matrix_damage_rows = get_risk_matrix_damage_rows(conn)
        top_scenarios_rows = get_top_scenarios_by_hazard_component(conn)
        data_substances = get_substances_by_component(conn)

        # Сводная таблица рисков гибели по составляющим (индивидуальный + коллективный)
        ind_map = {r.get("hazard_component"): r.get("individual_risk_fatalities") for r in individual_risk_rows}
        coll_map = {r.get("hazard_component"): r.get("collective_risk_fatalities") for r in collective_risk_rows}
        components = sorted({*ind_map.keys(), *coll_map.keys()}, key=lambda x: str(x))
        fatality_risk_by_component_rows = [
            {
                "hazard_component": comp,
                "individual_risk_fatalities": ind_map.get(comp),
                "collective_risk_fatalities": coll_map.get(comp),
            }
            for comp in components
        ]

    doc = Document(TEMPLATE_PATH)

    render_section_at_marker(
        doc=doc,
        marker="{{SUBSTANCES_SECTION}}",
        section_title="Сведения о веществах",
        items=substances,
        item_title_field="name",
        sections=SUBSTANCE_SECTIONS,
        json_formatter=pretty_json_substance,
    )

    render_section_at_marker(
        doc=doc,
        marker="{{EQUIPMENT_SECTION}}",
        section_title="Сведения об оборудовании",
        items=equipment,
        item_title_field="equipment_name",
        sections=EQUIPMENT_SECTIONS,
        json_formatter=pretty_json_generic,
    )

    render_distribution_table_at_marker(
        doc=doc,
        marker="{{DISTRIBUTION_SECTION}}",
        title="Распределение опасного вещества по оборудованию",
        rows=distribution,
    )

    render_scenarios_table_at_marker(
        doc=doc,
        marker="{{SCENARIOS_SECTION}}",
        title="Сценарии аварий",
        rows=scenarios,
    )

    render_ov_amount_table_at_marker(
        doc=doc,
        marker="{{OV_AMOUNT_SECTION}}",
        title="Оценка количества опасного вещества в аварии",
        rows=ov_amounts,
    )

    render_impact_zones_table(
        doc,
        "{{IMPACT_ZONES_SECTION}}",
        impact_zones
    )

    render_personnel_casualties_table_at_marker(
        doc=doc,
        marker="{{CASUALTIES_SECTION}}",
        title="Оценка количества погибших/пострадавших",
        rows=casualties,
    )

    render_damage_table_at_marker(
        doc=doc,
        marker="{{DAMAGE_SECTION}}",
        rows=damage_rows,
    )

    render_collective_risk_table(
        doc=doc,
        marker="{{COLLECTIVE_RISK_SECTION}}",
        rows=collective_risk_rows,
    )

    render_individual_risk_table(
        doc=doc,
        marker="{{INDIVIDUAL_RISK_SECTION}}",
        rows=individual_risk_rows,
    )

    render_fatal_accident_frequency_text(
        doc=doc,
        marker="{{FATAL_ACCIDENT_FREQUENCY}}",
        min_freq=min_f,
        max_freq=max_f,
    )

    render_max_damage_by_component_table(
        doc=doc,
        marker="{{MAX_DAMAGE_BY_COMPONENT_SECTION}}",
        rows=max_damage_rows,
    )

    render_top_scenarios_by_component_table(
        doc=doc,
        marker="{{TOP_SCENARIOS_BY_COMPONENT_SECTION}}",
        rows=top_scenarios_rows,
    )

    render_fatality_risk_by_component_table(
        doc=doc,
        marker="{{FATALITY_RISK_BY_COMPONENT_SECTION}}",
        rows=fatality_risk_by_component_rows,
    )

    render_comparative_fatality_risk_table(
        doc=doc,
        marker="{{COMPARATIVE_FATALITY_RISK_TABLE}}",
        individual_risk_rows=individual_risk_rows,
    )

    render_ngk_background_comparison_table(
        doc=doc,
        marker="{{NGK_BACKGROUND_RISK_COMPARISON}}",
        conn=conn,
    )

    render_substances_by_component_table(
        doc=doc,
        marker="{{SUBSTANCES_BY_COMPONENT_TABLE}}",
        conn=conn,
    )

    render_top_scenarios_description_by_component_table(
        doc=doc,
        marker="{{TOP_SCENARIOS_DESC_BY_COMPONENT}}",
        conn=conn,
    )

    # --- диаграммы ---
    render_fn_chart_at_marker(
        doc,
        "{{FN_CHART}}",
        fn_rows,
    )

    render_fg_chart_at_marker(
        doc,
        "{{FG_CHART}}",
        fg_rows,
    )

    render_pareto_damage_chart_at_marker(
        doc,
        "{{PARETO_DAMAGE_CHART}}",
        pareto_damage_rows,
    )

    render_pareto_fatalities_chart_at_marker(doc, "{{PARETO_FATALITIES_CHART}}", pareto_rows)
    render_pareto_injured_chart_at_marker(doc, "{{PARETO_INJURED_CHART}}", pareto_rows)

    render_pareto_environmental_damage_chart_at_marker(
        doc,
        "{{PARETO_ENV_DAMAGE_CHART}}",
        pareto_env_rows,
    )

    render_component_damage_chart_at_marker(
        doc,
        "{{DAMAGE_BY_COMPONENT_CHART}}",
        component_damage_rows,
    )

    render_risk_matrix_chart_at_marker(
        doc,
        "{{RISK_MATRIX_CHART}}",
        risk_matrix_rows,
    )

    render_risk_matrix_damage_chart_at_marker(
        doc,
        "{{RISK_MATRIX_DAMAGE_CHART}}",
        risk_matrix_damage_rows,
    )

    doc.save(OUT_PATH)
    print("Отчёт сформирован:", OUT_PATH)


if __name__ == "__main__":
    main()
