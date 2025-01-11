# report_generator.py
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io
from docx.enum.section import WD_ORIENT

from database.db_connection import DatabaseConnection
from database.repositories.calculation_repo import CalculationResultRepository
from models.equipment import EquipmentType
from models.substance import SubstanceType


class ReportGenerator:
    """Класс для генерации отчетов в Word"""

    def __init__(self, db: DatabaseConnection):
        self.db = db
        self.calc_repo = CalculationResultRepository(db)
        self.doc = Document()

    def _sum_data_for_fn(self, data: list):
        '''
        Функция вычисления суммирования вероятностей F при которой пострадало не менее N человек
        :param data: данные вида [[3.8e-08, 1],[5.8e-08, 2],[1.1e-08, 1]..]
        :return: данные вида: {1: 0.00018, 2: 0.012, 3: 6.9008e-06, 4: 3.8e-08, 5: 7.29e-05}
        '''
        uniq = set(sorted([i[1] for i in data]))
        result = dict(zip(uniq, [0] * len(uniq)))

        for item_data in data:
            for item_uniq in uniq:
                if item_data[1] >= item_uniq:
                    result[item_uniq] = result[item_uniq] + item_data[0]

        if 0 in result:
            del result[0]  # удалить суммарную вероятность где пострадало 0 человек
        return result

    def _sum_data_for_fg(self, data: list):
        '''
        Функция вычисления суммирования вероятностей F при которой ущерб не менее G млн.руб
        :param data: данные вида [[3.8e-08, 1.2],[5.8e-08, 0.2],[1.1e-08, 12.4]..]
        :return: данные вида: {0.2: 0.00018, 1: 0.012, 3: 6.9008e-06, 5: 3.8e-08, 6.25: 7.29e-05}
        '''
        uniq = np.arange(0, max([i[1] for i in data])+max([i[1] for i in data]) / 7, max([i[1] for i in data]) / 7)

        result = dict(zip(uniq, [0] * len(uniq)))

        for item_data in data:
            for item_uniq in uniq:
                if item_data[1] >= item_uniq:
                    result[item_uniq] = result[item_uniq] + item_data[0]

        del result[0]  # удалить суммарную вероятность где ущерб 0
        return result

    def add_critical_scenarios_table(self, project_code: str = None):
        """Добавление таблицы наиболее опасных и вероятных сценариев"""
        self.doc.add_heading('Наиболее опасные и наиболее вероятные сценарии', level=2)

        # Получаем результаты расчетов
        results = self.calc_repo.get_by_project(project_code) if project_code else self.calc_repo.get_all()
        if not results:
            self.doc.add_paragraph('Нет данных для отображения')
            return

        # Получаем уникальные компоненты
        components = set()
        for result in results:
            component = None
            if hasattr(result, 'component_enterprise') and result.component_enterprise:
                component = result.component_enterprise
            elif hasattr(result, 'equipment_name') and result.equipment_name:
                import re
                match = re.search(r'\((.*?)\)', result.equipment_name)
                if match:
                    component = match.group(1)
            if component:
                components.add(component)

        # Создаем таблицу
        headers = [
            "Составляющая",
            "Тип сценария",
            "№ сценария",
            "Оборудование",
            "Погибшие",
            "Пострадавшие",
            "Суммарный ущерб",
            "Частота, 1/год"
        ]

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Заполняем заголовки
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Для каждого компонента находим наиболее опасный и вероятный сценарии
        for component in sorted(components):
            # Фильтруем результаты для текущего компонента
            component_results = []
            for result in results:
                if (hasattr(result, 'component_enterprise') and
                        result.component_enterprise == component):
                    component_results.append(result)
                elif hasattr(result, 'equipment_name'):
                    match = re.search(r'\((.*?)\)', result.equipment_name)
                    if match and match.group(1) == component:
                        component_results.append(result)

            if not component_results:
                continue

            # Находим наиболее опасный сценарий (по суммарному ущербу)
            most_dangerous = max(component_results, key=lambda x: x.total_damage)

            # Добавляем строку для наиболее опасного сценария
            row_cells = table.add_row().cells
            row_cells[0].text = component
            row_cells[1].text = "Наиболее опасный"
            row_cells[2].text = str(most_dangerous.scenario_number)
            row_cells[3].text = most_dangerous.equipment_name
            row_cells[4].text = str(most_dangerous.casualties)
            row_cells[5].text = str(most_dangerous.injured)
            row_cells[6].text = f"{most_dangerous.total_damage:.2f}"
            row_cells[7].text = f"{most_dangerous.probability:.2e}"

            # Находим наиболее вероятный сценарий
            most_probable = max(component_results, key=lambda x: x.probability)

            # Добавляем строку для наиболее вероятного сценария
            row_cells = table.add_row().cells
            row_cells[0].text = component
            row_cells[1].text = "Наиболее вероятный"
            row_cells[2].text = str(most_probable.scenario_number)
            row_cells[3].text = most_probable.equipment_name
            row_cells[4].text = str(most_probable.casualties)
            row_cells[5].text = str(most_probable.injured)
            row_cells[6].text = f"{most_probable.total_damage:.2f}"
            row_cells[7].text = f"{most_probable.probability:.2e}"

        # Форматируем таблицу
        for row in table.rows[1:]:  # Пропускаем заголовок
            for cell in row.cells[2:]:  # Центрируем все колонки кроме первых двух
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_risk_component_analysis(self, project_code: str = None):
        """Добавление таблицы анализа риска по компонентам"""
        self.doc.add_heading('Анализ риска по компонентам', level=2)

        # Получаем результаты расчетов
        results = self.calc_repo.get_by_project(project_code) if project_code else self.calc_repo.get_all()
        if not results:
            self.doc.add_paragraph('Нет данных для отображения')
            return

        # Создаем таблицу
        headers = [
            "№ п/п",
            "Составляющая",
            "Макс. ущерб, млн.руб",
            "Макс. экологический ущерб, млн.руб",
            "Макс. количество погибших, чел",
            "Макс. количество пострадавших, чел",
            "Коллективный риск гибели, чел/год",
            "Коллективный риск ранения, чел/год",
            "Индивидуальный риск гибели, чел/год",
            "Индивидуальный риск ранения, чел/год",
            "Уровень риска, ppm",
            "Уровень риска, дБR",
            "Частота аварии, 1/год"
        ]

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Заполняем заголовки
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Получаем уникальные компоненты
        components = set()
        for result in results:
            component = None
            if hasattr(result, 'component_enterprise') and result.component_enterprise:
                component = result.component_enterprise
            elif hasattr(result, 'equipment_name') and result.equipment_name:
                import re
                match = re.search(r'\((.*?)\)', result.equipment_name)
                if match:
                    component = match.group(1)
            if component:
                components.add(component)

        # Заполняем данные
        from models.risk_analysis import ComponentRiskAnalysis
        from models.dangerous_object import DangerousObject

        dangerous_object = None  # здесь можно добавить получение ОПО если нужно

        for i, component in enumerate(sorted(components), 1):
            # Расчет анализа риска для компонента
            analysis = ComponentRiskAnalysis.calculate_for_component(
                component, results, dangerous_object
            )

            if not analysis:
                continue

            # Добавляем строку
            row_cells = table.add_row().cells

            # Заполняем ячейки
            row_cells[0].text = str(i)
            row_cells[1].text = analysis.component_name
            row_cells[2].text = f"{analysis.max_damage:.2f}"
            row_cells[3].text = f"{analysis.max_eco_damage:.2f}"
            row_cells[4].text = str(analysis.max_casualties)
            row_cells[5].text = str(analysis.max_injured)
            row_cells[6].text = f"{analysis.collective_death_risk:.2e}"
            row_cells[7].text = f"{analysis.collective_injury_risk:.2e}"
            row_cells[8].text = f"{analysis.individual_death_risk:.2e}"
            row_cells[9].text = f"{analysis.individual_injury_risk:.2e}"
            row_cells[10].text = f"{analysis.risk_level_ppm:.2f}"
            row_cells[11].text = f"{analysis.risk_level_dbr:.2f}"
            row_cells[12].text = f"{analysis.max_death_frequency:.2e}"

    def setup_page_format(self):
        """Настройка формата страницы А3 горизонтальный"""
        section = self.doc.sections[0]

        # Размер А3 в дюймах (297x420 мм)
        # 1 дюйм = 914400 твипов
        section.page_height = int(11.69 * 914400)  # 297мм
        section.page_width = int(16.54 * 914400)  # 420мм

        # Устанавливаем горизонтальную ориентацию
        section.orientation = WD_ORIENT.LANDSCAPE

        # Поля страницы (1 дюйм = 914400 твипов)
        margin = int(0.8 * 914400)  # 20мм
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin

    def generate_full_report(self, file_path: str, project_code: str = None):
        """Генерация полного отчета"""
        # Настраиваем формат страницы
        self.setup_page_format()

        # Заголовок документа
        self.doc.add_heading('Отчет по результатам расчетов', 0)
        # Добавляем новые таблицы в начало отчета
        self.add_organization_info(project_code)
        self.doc.add_page_break()

        self.add_opo_info(project_code)
        self.doc.add_page_break()

        self.add_project_info(project_code)
        self.doc.add_page_break()

        # Добавляем таблицу с описанием оборудования
        self.add_equipment_table(project_code)
        # Явно добавляем разрыв страницы
        self.doc.add_page_break()

        # Добавляем таблицу распределения вещества
        self.doc.add_page_break()
        self.add_substance_distribution_table(project_code)

        # Таблица результатов расчета
        self.add_calculation_results_table(project_code)

        # Явно добавляем разрыв страницы перед анализом риска
        self.doc.add_page_break()

        # Добавляем раздел анализа риска
        self.add_risk_analysis(project_code)

        # Добавляем новую таблицу критических сценариев
        self.add_critical_scenarios_table(project_code)

        # Сохраняем документ
        self.doc.save(file_path)

    def add_risk_analysis(self, project_code: str = None):
        """Добавление раздела анализа риска"""
        self.doc.add_heading('Анализ риска', level=1)

        # Добавляем статистику
        self.add_risk_statistics(project_code)

        # Добавляем таблицу анализа по компонентам
        self.add_risk_component_analysis(project_code)

        # Явно добавляем разрыв страницы перед диаграммами
        self.doc.add_page_break()

        # Добавляем F/N и F/G диаграммы
        self.add_fn_fg_diagrams(project_code)

    def add_equipment_table(self, project_code: str = None):
        """Добавление таблицы с описанием оборудования"""
        self.doc.add_heading('Характеристика оборудования', level=1)

        # Получаем все оборудование для проекта
        query = """
        SELECT 
            be.id,
            be.name,
            be.equipment_type,
            be.component_enterprise,
            be.pressure,
            be.temperature,
            -- Данные трубопроводов
            p.diameter_category,
            p.length_meters,
            p.diameter_pipeline,
            p.flow as pipeline_flow,
            -- Данные насосов
            pm.pump_type,
            pm.volume as pump_volume,
            pm.flow as pump_flow,
            -- Данные тех. устройств
            td.device_type,
            td.volume as device_volume,
            td.degree_filling as device_filling,
            -- Данные резервуаров
            t.tank_type,
            t.volume as tank_volume,
            t.degree_filling as tank_filling,
            -- Данные автоцистерн
            tt.pressure_type,
            tt.volume as truck_volume,
            tt.degree_filling as truck_filling,
            -- Данные компрессоров
            c.comp_type,
            c.volume as comp_volume,
            c.flow as comp_flow
        FROM base_equipment be
        LEFT JOIN pipelines p ON p.id = be.id
        LEFT JOIN pumps pm ON pm.id = be.id
        LEFT JOIN technological_devices td ON td.id = be.id
        LEFT JOIN tanks t ON t.id = be.id
        LEFT JOIN truck_tanks tt ON tt.id = be.id
        LEFT JOIN compressors c ON c.id = be.id
        WHERE be.project_id IN (SELECT id FROM projects WHERE project_code = ?)
        ORDER BY be.name
        """

        equipment_data = self.db.execute_query(query, (project_code,)) if project_code else []

        if not equipment_data:
            self.doc.add_paragraph('Нет данных для отображения')
            return

        # Создаем таблицу
        headers = [
            "Составляющая",
            "Наименование оборудования",
            "Расположение",
            "Кол-во, шт.",
            "Назначение",
            "Техническая характеристика"
        ]

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Заполняем заголовки
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Группируем оборудование по составляющим
        components_dict = {}

        for eq in equipment_data:
            component = eq['component_enterprise']
            if component not in components_dict:
                components_dict[component] = []

            # Формируем технические характеристики в зависимости от типа оборудования
            tech_specs = []
            eq_type = eq['equipment_type']

            # Общие характеристики
            tech_specs.append(f"P={eq['pressure']:.2f} МПа")
            tech_specs.append(f"T={eq['temperature']:.1f} °C")

            if eq_type == 'Pipeline':
                tech_specs.append(f"L={eq['length_meters']:.1f} м")
                tech_specs.append(f"D={eq['diameter_pipeline']:.1f} мм")
                if eq['pipeline_flow']:
                    tech_specs.append(f"Q={eq['pipeline_flow']:.2f} кг/с")
                # tech_specs.append(f"Категория: {eq['diameter_category']}")

            elif eq_type == 'Pump':
                if eq['pump_volume']:
                    tech_specs.append(f"V={eq['pump_volume']:.1f} м³")
                if eq['pump_flow']:
                    tech_specs.append(f"Q={eq['pump_flow']:.2f} кг/с")
                # tech_specs.append(f"Тип: {eq['pump_type']}")

            elif eq_type == 'Technological_device':
                if eq['device_volume']:
                    tech_specs.append(f"V={eq['device_volume']:.1f} м³")
                if eq['device_filling']:
                    tech_specs.append(f"Степень заполнения={eq['device_filling'] * 100:.1f}%")
                # tech_specs.append(f"Тип: {eq['device_type']}")

            elif eq_type == 'Tank':
                if eq['tank_volume']:
                    tech_specs.append(f"V={eq['tank_volume']:.1f} м³")
                if eq['tank_filling']:
                    tech_specs.append(f"Степень заполнения={eq['tank_filling'] * 100:.1f}%")
                # tech_specs.append(f"Тип: {eq['tank_type']}")

            elif eq_type == 'Truck_tank':
                if eq['truck_volume']:
                    tech_specs.append(f"V={eq['truck_volume']:.1f} м³")
                if eq['truck_filling']:
                    tech_specs.append(f"Степень заполнения={eq['truck_filling'] * 100:.1f}%")
                # tech_specs.append(f"Тип: {eq['pressure_type']}")

            elif eq_type == 'Compressor':
                if eq['comp_volume']:
                    tech_specs.append(f"V={eq['comp_volume']:.1f} м³")
                if eq['comp_flow']:
                    tech_specs.append(f"Q={eq['comp_flow']:.2f} кг/с")
                # tech_specs.append(f"Тип: {eq['comp_type']}")

            components_dict[component].append({
                'name': eq['name'],
                'type': eq_type,
                'tech_specs': "\n ".join(tech_specs)
            })

        # Заполняем таблицу
        for component, equipments in components_dict.items():
            for equipment in sorted(equipments, key=lambda x: x['name']):
                row_cells = table.add_row().cells

                # Составляющая
                row_cells[0].text = component

                # Наименование оборудования
                row_cells[1].text = equipment['name'].split('(')[0].strip()

                # Расположение
                row_cells[2].text = "Наземное"

                # Количество
                row_cells[3].text = "1"

                # Назначение
                if equipment['type'] == 'Pipeline':
                    row_cells[4].text = "Транспортировка опасного вещества"
                elif equipment['type'] == 'Pump':
                    row_cells[4].text = "Перекачивание опасного вещества"
                elif equipment['type'] in ['Tank', 'Truck_tank']:
                    row_cells[4].text = "Хранение опасного вещества"
                elif equipment['type'] == 'Technological_device':
                    row_cells[4].text = "Проведение технологического процесса"
                elif equipment['type'] == 'Compressor':
                    row_cells[4].text = "Компримирование опасного вещества"

                # Техническая характеристика
                row_cells[5].text = equipment['tech_specs']

        # Применяем форматирование
        for row in table.rows:
            for cell in row.cells:
                if cell != row.cells[0] and cell != row.cells[1] and cell != row.cells[5]:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_substance_distribution_table(self, project_code: str = None):
        """Добавление таблицы распределения опасного вещества по оборудованию"""
        self.doc.add_heading('Распределение опасного вещества по оборудованию', level=1)

        if not project_code:
            self.doc.add_paragraph('Не указан код проекта')
            return

        # Сначала получим все оборудование проекта
        base_query = """
        SELECT 
            be.id,
            be.name,
            be.equipment_type,
            be.component_enterprise,
            be.temperature,
            p.diameter_category,
            p.length_meters,
            p.diameter_pipeline,
            p.flow as pipeline_flow,
            pm.pump_type,
            pm.volume as pump_volume,
            pm.flow as pump_flow,
            td.device_type,
            td.volume as device_volume,
            td.degree_filling as device_filling,
            t.tank_type,
            t.volume as tank_volume,
            t.degree_filling as tank_filling,
            tt.pressure_type,
            tt.volume as truck_volume,
            tt.degree_filling as truck_filling,
            c.comp_type,
            c.volume as comp_volume,
            c.flow as comp_flow,
            (SELECT cr.mass_in_equipment 
             FROM calculation_results cr 
             WHERE cr.project_code = ? 
             AND cr.equipment_name = be.name 
             ORDER BY cr.scenario_number DESC 
             LIMIT 1) as mass_in_equipment
        FROM base_equipment be
        LEFT JOIN pipelines p ON p.id = be.id
        LEFT JOIN pumps pm ON pm.id = be.id
        LEFT JOIN technological_devices td ON td.id = be.id
        LEFT JOIN tanks t ON t.id = be.id
        LEFT JOIN truck_tanks tt ON tt.id = be.id
        LEFT JOIN compressors c ON c.id = be.id
        WHERE be.project_id IN (SELECT id FROM projects WHERE project_code = ?)
        ORDER BY be.name
        """

        equipment_data = self.db.execute_query(base_query, (project_code, project_code))

        if not equipment_data:
            self.doc.add_paragraph('Нет данных для отображения')
            return

        # Создаем таблицу
        headers = [
            "Составляющая",
            "Наименование оборудования",
            "Кол-во единиц",
            "В единице оборудования, т",
            "В блоке, т",
            "Агр. состояние",
            "Температура, °C"
        ]

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Заполняем заголовки
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Группируем оборудование по составляющим
        components_dict = {}

        for eq in equipment_data:
            component = eq['component_enterprise']
            if not component:  # Пропускаем записи без компонента
                continue

            if component not in components_dict:
                components_dict[component] = []

            # Определяем агрегатное состояние
            if eq['equipment_type'] == 'Compressor':
                agr_state = "г.ф."
            else:
                agr_state = "ж.ф."

            # Используем значение mass_in_equipment из результатов расчета, если оно есть
            mass = eq['mass_in_equipment'] if eq['mass_in_equipment'] is not None else 0

            components_dict[component].append({
                'name': eq['name'],
                'mass': mass,
                'temperature': eq['temperature'],
                'agr_state': agr_state
            })

        # Заполняем таблицу
        for component, equipments in components_dict.items():
            for equipment in sorted(equipments, key=lambda x: x['name']):
                row_cells = table.add_row().cells

                # Составляющая
                row_cells[0].text = component

                # Наименование оборудования - берем только часть до скобки
                name_parts = equipment['name'].split('(')
                equipment_name = name_parts[0].strip()
                row_cells[1].text = equipment_name

                # Количество единиц
                row_cells[2].text = "1"

                # В единице оборудования
                row_cells[3].text = f"{equipment['mass']:.3f}"

                # В блоке
                row_cells[4].text = f"{equipment['mass']:.3f}"

                # Агрегатное состояние
                row_cells[5].text = equipment['agr_state']

                # Температура
                row_cells[6].text = f"{equipment['temperature']:.1f}"

                # Центрируем все ячейки кроме первых двух
                for cell in row_cells[2:]:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_calculation_results_table(self, project_code: str = None):
        """Добавление таблицы результатов расчета"""
        self.doc.add_heading('Результаты расчета', level=1)

        # Получаем результаты
        if project_code:
            results = self.calc_repo.get_by_project(project_code)
        else:
            results = self.calc_repo.get_all()

        if not results:
            self.doc.add_paragraph('Нет данных для отображения')
            return
        """Добавление подробной таблицы результатов"""

        # Получаем все результаты
        results = self.calc_repo.get_by_project(project_code) if project_code else self.calc_repo.get_all()
        # Сортируем результаты по номеру сценария
        results = sorted(results, key=lambda x: int(x.scenario_number))

        if not results:
            self.doc.add_paragraph('Нет данных для отображения')
            return

        # Создаем таблицу
        headers = [
            "№ п/п", "Код проекта", "№ сценария", "Оборудование",
            "Тип оборудования", "Тип вещества",
            "q_10.5 (кВт/м2)", "q_7.0 (кВт/м2)", "q_4.2 (кВт/м2)", "q_1.4 (кВт/м2)",
            "p_53 (кПа)", "p_28 (кПа)", "p_12 (кПа)", "p_5 (кПа)", "p_2 (кПа)",
            "Длина факела (м)", "Диаметр факела (м)",
            "Радиус НКПР (м)", "Радиус вспышки (м)",
            "Погибшие (чел)", "Пострадавшие (чел)",
            "Прямые потери (млн.руб)", "Затраты на ЛЛА (млн.руб)",
            "Социальный ущерб (млн.руб)", "Косвенный ущерб (млн.руб)",
            "Экологический ущерб (млн.руб)", "Суммарный ущерб (млн.руб)",
            "Риск гибели (чел/год)", "Риск травмирования (чел/год)",
            "Ожидаемый ущерб (млн.руб/год)", "Вероятность (1/год)",
            "Масса в аварии (т)", "Масса в ПФ (т)", "Масса в оборудовании (т)"
        ]

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Заполняем заголовки
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


        # Заполняем данные таблицы
        for i, result in enumerate(results, 1):
            row_cells = table.add_row().cells

            # Базовая информация
            row_cells[0].text = str(i)
            row_cells[1].text = result.project_code
            row_cells[2].text = str(result.scenario_number)
            row_cells[3].text = result.equipment_name
            row_cells[4].text = EquipmentType.get_display_name(result.equipment_type)
            row_cells[5].text = SubstanceType.get_display_name(result.substance_type)

            # Параметры теплового излучения
            row_cells[6].text = f"{result.q_10_5:.2f}" if result.q_10_5 else "-"
            row_cells[7].text = f"{result.q_7_0:.2f}" if result.q_7_0 else "-"
            row_cells[8].text = f"{result.q_4_2:.2f}" if result.q_4_2 else "-"
            row_cells[9].text = f"{result.q_1_4:.2f}" if result.q_1_4 else "-"

            # Параметры ударной волны
            row_cells[10].text = f"{result.p_53:.2f}" if result.p_53 else "-"
            row_cells[11].text = f"{result.p_28:.2f}" if result.p_28 else "-"
            row_cells[12].text = f"{result.p_12:.2f}" if result.p_12 else "-"
            row_cells[13].text = f"{result.p_5:.2f}" if result.p_5 else "-"
            row_cells[14].text = f"{result.p_2:.2f}" if result.p_2 else "-"

            # Параметры факела
            row_cells[15].text = f"{result.l_f:.2f}" if result.l_f else "-"
            row_cells[16].text = f"{result.d_f:.2f}" if result.d_f else "-"

            # Параметры вспышки
            row_cells[17].text = f"{result.r_nkpr:.2f}" if result.r_nkpr else "-"
            row_cells[18].text = f"{result.r_flash:.2f}" if result.r_flash else "-"

            # Последствия
            row_cells[19].text = str(result.casualties)
            row_cells[20].text = str(result.injured)

            # Ущерб
            row_cells[21].text = f"{result.direct_losses:.2f}"
            row_cells[22].text = f"{result.liquidation_costs:.2f}"
            row_cells[23].text = f"{result.social_losses:.2f}"
            row_cells[24].text = f"{result.indirect_damage:.2f}"
            row_cells[25].text = f"{result.environmental_damage:.2f}"
            row_cells[26].text = f"{result.total_damage:.2f}"

            # Риски
            row_cells[27].text = f"{result.casualty_risk:.2e}"
            row_cells[28].text = f"{result.injury_risk:.2e}"
            row_cells[29].text = f"{result.expected_damage:.2e}"
            row_cells[30].text = f"{result.probability:.2e}"

            # Массы
            row_cells[31].text = f"{result.mass_in_accident:.2f}"
            row_cells[32].text = f"{result.mass_in_factor:.2f}"
            row_cells[33].text = f"{result.mass_in_equipment:.2f}"

    # report_generator.py (продолжение)

    def add_risk_statistics(self, project_code: str = None):
        """Добавление статистических показателей"""
        results = self.calc_repo.get_by_project(project_code) if project_code else self.calc_repo.get_all()

        self.doc.add_heading('Статистические показатели', level=2)

        stats_table = self.doc.add_table(rows=8, cols=2)
        stats_table.style = 'Table Grid'

        # Заполняем статистику
        rows = stats_table.rows
        rows[0].cells[0].text = "Всего сценариев"
        rows[0].cells[1].text = str(len(results))

        rows[1].cells[0].text = "Максимальное число погибших"
        rows[1].cells[1].text = str(max(r.casualties for r in results))

        rows[2].cells[0].text = "Максимальное число пострадавших"
        rows[2].cells[1].text = str(max(r.injured for r in results))

        rows[3].cells[0].text = "Максимальный ущерб (млн.руб)"
        rows[3].cells[1].text = f"{max(r.total_damage for r in results):.2f}"

        rows[4].cells[0].text = "Суммарный риск гибели (чел/год)"
        rows[4].cells[1].text = f"{sum(r.casualty_risk for r in results):.2e}"

        rows[5].cells[0].text = "Суммарный риск травмирования (чел/год)"
        rows[5].cells[1].text = f"{sum(r.injury_risk for r in results):.2e}"

        rows[6].cells[0].text = "Максимальная частота аварий с гибелью (1/год)"
        death_scenarios = [r for r in results if r.casualties >= 1]
        max_death_frequency = max((r.probability for r in death_scenarios), default=0)
        rows[6].cells[1].text = f"{max_death_frequency:.2e}"

        rows[7].cells[0].text = "Максимальный экологический ущерб (млн.руб)"
        rows[7].cells[1].text = f"{max(r.environmental_damage for r in results):.2f}"

    # report_generator.py (продолжение)

    def add_fn_fg_diagrams(self, project_code: str = None):
        """Добавление F/N и F/G диаграмм"""
        self.doc.add_heading('F/N и F/G диаграммы', level=2)

        results = self.calc_repo.get_by_project(project_code) if project_code else self.calc_repo.get_all()

        # Создаем F/N диаграмму
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 5))

        # F/N диаграмма
        # Построение Fn диаграммы
        casualty_data = []
        for result in results:
            if result.casualties > 0:
                casualty_data.append((result.probability, result.casualties))

        sum_data = self._sum_data_for_fn(casualty_data)

        if casualty_data:
            people, probability = list(sum_data.keys()), list(sum_data.values())
            # для сплошных горизонтальных линий
            chart_line_x = []
            chart_line_y = []

            for i in range(len(people) - 1):
                # Добавляем две точки для текущего горизонтального отрезка
                chart_line_x.extend([people[i], people[i + 1], None])
                chart_line_y.extend([probability[i], probability[i], None])

            print(chart_line_x, chart_line_y)
            # для вертикальных пунктирных линий
            chart_dot_line_x = []
            chart_dot_line_y = []

            for i in range(len(people) - 1):
                # Добавляем две точки для вертикального пунктирного отрезка
                chart_dot_line_x.extend([people[i + 1], people[i + 1]])
                chart_dot_line_y.extend([probability[i], probability[i + 1]])

            # Строим F/N кривую
            ax1.semilogy(chart_line_x, chart_line_y, color='b', linestyle='-', marker='.')
            ax1.semilogy(chart_dot_line_x, chart_dot_line_y, color='b', linestyle='--', marker='.')
            ax1.set_xticks(people)
            ax1.grid(True)
            ax1.set_xlabel('N, число погибших')
            ax1.set_ylabel('F, частота событий с N и более погибшими, 1/год')
            ax1.set_title('F/N диаграмма')

        # F/G диаграмма
        damage_data = []
        for result in results:
            if result.casualties > 0:
                damage_data.append((result.probability, result.casualties))

        sum_data = self._sum_data_for_fg(damage_data)

        if damage_data:
            damage, probability = list(sum_data.keys()), list(sum_data.values())
            # для сплошных линий
            chart_line_x = []
            chart_line_y = []
            for i in damage:
                if damage[0] == i:
                    chart_line_x.extend([0, i, i, i])
                    chart_line_y.extend([probability[damage.index(i)], probability[damage.index(i)], None, None])
                elif damage[-1] == i:
                    chart_line_x.extend([damage[damage.index(i)-1], damage[damage.index(i)-1], i, i])
                    chart_line_y.extend([probability[damage.index(i)], probability[damage.index(i)], probability[damage.index(i)], probability[damage.index(i)]])
                    break
                else:
                    chart_line_x.extend([damage[damage.index(i) - 1], i, i, i])
                    chart_line_y.extend([probability[damage.index(i)], probability[damage.index(i)], None, None])

            # для пунктирных линий
            chart_dot_line_x = []
            chart_dot_line_y = []
            for i in damage:
                if i == damage[-1]:
                    chart_dot_line_x.extend([i, i])
                    chart_dot_line_y.extend([probability[damage.index(i)], probability[damage.index(i)]])
                    chart_dot_line_x.extend([i, i])
                    chart_dot_line_y.extend([probability[damage.index(i)], 0])
                    break
                chart_dot_line_x.extend([i, i])
                chart_dot_line_y.extend([probability[damage.index(i)], probability[damage.index(i) + 1]])

            # Создание графика
            # Построение основной диаграммы
            ax2.semilogy(chart_line_x, chart_line_y, color='r', linestyle='-', marker='.')
            ax2.semilogy(chart_dot_line_x, chart_dot_line_y, color='r', linestyle='--', marker='.')
            ax2.grid(True)
            ax2.set_xticks(damage)
            ax2.set_xlabel('G, ущерб, млн.руб')
            ax2.set_ylabel('F, частота событий с ущербом G и более, 1/год')
            ax2.set_title('F/G диаграмма')

        plt.tight_layout()

        # Сохраняем диаграммы во временный буфер
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)

        # Добавляем изображение в документ
        self.doc.add_picture(img_buffer, width=Inches(7))

        # Закрываем графики
        plt.close()


    def add_project_info(self, project_code: str = None):
        """Добавление информации о проекте"""
        self.doc.add_heading('Информация о проекте', level=1)

        if not project_code:
            self.doc.add_paragraph('Код проекта не указан')
            return

        # Получаем данные проекта
        query = "SELECT * FROM projects WHERE project_code = ?"
        result = self.db.execute_query(query, (project_code,))

        if not result:
            self.doc.add_paragraph('Проект не найден')
            return

        project_data = result[0]

        # Создаем таблицу
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Заполняем данные
        rows = [
            ("Наименование проекта", project_data['name']),
            ("Код проекта", project_data['project_code']),
            ("Описание", project_data['description'] or '-'),
            ("Описание автоматизации", project_data['automation_description'] or '-'),
            ("Шифр ДПБ", project_data['dpb_code'] or '-'),
            ("Шифр РПЗ", project_data['rpz_code'] or '-'),
            ("Шифр ИФЛ", project_data['ifl_code'] or '-'),
            ("Шифр ГОЧС", project_data['gochs_code'] or '-'),
            ("Шифр МПБ", project_data['mpb_code'] or '-')
        ]

        for row_data in rows:
            cells = table.add_row().cells
            cells[0].text = row_data[0]
            cells[1].text = str(row_data[1])


    def add_opo_info(self, project_code: str = None):
        """Добавление информации об ОПО"""
        self.doc.add_heading('Информация об опасном производственном объекте', level=1)

        if not project_code:
            return

        # Получаем ОПО через связь с проектом
        query = """
            SELECT do.* FROM dangerous_objects do
            JOIN projects p ON p.opo_id = do.id
            WHERE p.project_code = ?
        """
        result = self.db.execute_query(query, (project_code,))

        if not result:
            self.doc.add_paragraph('ОПО не найден')
            return

        opo_data = result[0]

        # Создаем таблицу
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Заполняем данные
        from models.dangerous_object import HazardClass
        rows = [
            ("Наименование ОПО", opo_data['name']),
            ("Регистрационный номер", opo_data['reg_number']),
            ("Класс опасности", HazardClass.get_display_name(HazardClass(opo_data['hazard_class']))),
            ("Место расположения", opo_data['location']),
            ("Количество работников", str(opo_data['employee_count'])),
            ("Класс загроможденности", f"Класс {opo_data['view_space']}")
        ]

        for row_data in rows:
            cells = table.add_row().cells
            cells[0].text = row_data[0]
            cells[1].text = row_data[1]

    def add_organization_info(self, project_code: str = None):
        """Добавление информации об организации"""
        self.doc.add_heading('Информация об организации', level=1)

        if not project_code:
            return

        # Получаем организацию через связь проект -> ОПО -> организация
        query = """
            SELECT org.* FROM organizations org
            JOIN dangerous_objects do ON do.organization_id = org.id
            JOIN projects p ON p.opo_id = do.id
            WHERE p.project_code = ?
        """
        result = self.db.execute_query(query, (project_code,))

        if not result:
            self.doc.add_paragraph('Организация не найдена')
            return

        org_data = result[0]

        # Создаем таблицу
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Заполняем данные
        rows = [
            ("Наименование", org_data['name']),
            ("Полное наименование", org_data['full_name']),
            ("Форма собственности", org_data['org_form']),
            ("Должность руководителя", org_data['head_position'] or '-'),
            ("ФИО руководителя", org_data['head_name'] or '-'),
            ("Юридический адрес", org_data['legal_address'] or '-'),
            ("Телефон", org_data['phone'] or '-'),
            ("Факс", org_data['fax'] or '-'),
            ("Email", org_data['email'] or '-'),
            ("Номер лицензии", org_data['license_number'] or '-'),
            ("Дата лицензии", org_data['license_date'] or '-'),
            ("Система промышленной безопасности", org_data['ind_safety_system'] or '-'),
            ("Производственный контроль", org_data['prod_control'] or '-'),
            ("Порядок расследования аварий", org_data['accident_investigation'] or '-'),
            ("Договор с АСФ", org_data['rescue_contract'] or '-'),
            ("Свидетельство АСФ", org_data['rescue_certificate'] or '-'),
            ("Договор с ПСФ", org_data['fire_contract'] or '-'),
            ("Свидетельство НАСФ", org_data['emergency_certificate'] or '-'),
            ("Материальные резервы", org_data['material_reserves'] or '-'),
            ("Финансовые резервы", org_data['financial_reserves'] or '-')
        ]

        # Заголовок таблицы
        header_cells = table.rows[0].cells
        header_cells[0].text = "Параметр"
        header_cells[1].text = "Значение"

        # Заполняем данные
        for row_data in rows:
            cells = table.add_row().cells
            cells[0].text = row_data[0]
            cells[1].text = str(row_data[1])

            # Выравнивание в ячейках
            cells[0].paragraphs[0].alignment = 1  # По левому краю
            cells[1].paragraphs[0].alignment = 1  # По левому краю

        # Форматируем заголовок
        for cell in table.rows[0].cells:
            cell.paragraphs[0].alignment = 1  # По центру
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True