import json
from pathlib import Path
from docx.shared import Cm
from docx import Document

from report.reportgen.constants import NGK_BACKGROUND_RISK
from report.reportgen.formatters import risk_to_dbr, format_float_2
from report.reportgen.constants import SOCIAL_FATALITY_RISKS_DBR
from core.path import (
    DB_PATH,
    REPORT_TEMPLATE_DIR,
    REPORT_OUTPUT_DIR,
    TYPICAL_SCENARIOS_PATH,
    ORGANIZATION_PATH,
    ORGANIZATION_SITE_ID,
    PROJECT_COMMON_PATH,
)

TEMPLATE_DIR = REPORT_TEMPLATE_DIR

# алиасы для минимальных правок ниже по файлу

OUT_PATH = REPORT_OUTPUT_DIR / "template_report_out.docx"

from report.reportgen.db import (
    open_db,
    get_used_substances,
    get_all_substances,
    get_used_equipment,
    get_hazard_distribution,
    get_scenarios,
    get_ov_amounts_in_accident,
    get_impact_zones,
    get_personnel_casualties,
    get_damage,
    get_collective_risk,
    get_individual_risk,
    get_fatal_accident_frequency_range,
    get_max_damage_by_hazard_component,
    get_fn_source_rows,
    get_fg_source_rows,
    get_pareto_risk_source_rows,
    get_pareto_damage_source_rows,
    get_pareto_environmental_damage_source_rows,
    get_max_losses_by_hazard_component,
    get_risk_matrix_rows,
    get_risk_matrix_damage_rows,
    get_top_scenarios_by_hazard_component,
    get_damage_by_component,
    get_substances_by_component,
    get_calculation_row_for_top_scenario,
    get_fatalities_injured_for_top_scenario,
    get_total_damage_for_top_scenario,
    get_ov_in_accident_for_top_scenario,
)
from report.reportgen.sections import SUBSTANCE_SECTIONS, EQUIPMENT_SECTIONS, _format_pf_zones, _detect_method_text
from report.reportgen.formatters import (
    format_value,
    pretty_json_substance,
    pretty_json_generic,
    format_exp,
    format_float_3,
    format_float_1,
)
from report.reportgen.word_utils import (
    find_paragraph_with_marker,
    clear_paragraph,
    insert_paragraph_after,
    insert_table_after,
    insert_paragraph_after_table,
    add_section_header_row,
    set_run_font,
)

from report.reportgen.charts import (
    build_fn_points,
    build_fg_points,
    save_fn_chart,
    save_fg_chart,
    build_pareto_series,
    save_pareto_chart,
    save_component_damage_chart,
    save_risk_matrix_chart,
    save_risk_matrix_chart_damage,
)


def set_repeat_table_header(row):
    """Делает строку таблицы повторяемой шапкой на каждой странице."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = trPr.find(qn("w:tblHeader"))
    if tblHeader is None:
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)
    tblHeader.set(qn("w:val"), "1")


def load_project_common() -> dict:
    """Читает data/project_common.json или возвращает {}."""
    p = PROJECT_COMMON_PATH
    if not p.exists():
        return {}
    with open(p, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data if isinstance(data, dict) else {}


def iter_variant_templates(template_dir: Path) -> list[Path]:
    """Все .docx в папке варианта, стабильно отсортировано."""
    if not template_dir.exists():
        return []
    return sorted([p for p in template_dir.glob("*.docx") if p.is_file()])


def clear_output_docx(output_dir: Path):
    """Удаляет только .docx в output_dir. Папки (charts) и прочее не трогает."""
    if not output_dir.exists():
        return
    for p in output_dir.glob("*.docx"):
        if p.is_file():
            p.unlink()


def delete_paragraph(paragraph):
    """Удаляет paragraph из документа (чтобы не оставался пустой разрыв строки)."""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_table_autofit_to_contents(table):
    """
    Включает автоподбор ширины колонок по содержимому.
    Важно: снимает наши принудительные tblLayout=fixed и tblW=100% (pct),
    иначе Word часто игнорирует autofit.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr

    # убрать fixed layout
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is not None:
        tblPr.remove(tblLayout)

    # убрать принудительную ширину таблицы (pct 100%)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is not None:
        tblPr.remove(tblW)

    # включить autofit
    table.autofit = True


def set_table_full_width(doc: Document, table, cols: int = 2, left_ratio: float = 0.5):
    """
    Растягивает таблицу на ширину рабочей области страницы и фиксирует layout,
    чтобы Word реально применил ширину (аналог "Автоподбор по ширине окна").

    - Для cols==2: деление по left_ratio (например 0.35 => 35/65).
    - Для cols>2: равномерное деление ширины по всем колонкам.

    Важно: функция задаёт tblW=100% и tblLayout=fixed на уровне XML.
    """
    # 1) фиксированный layout (иначе Word может игнорировать widths)
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

    # 2) ширина таблицы = 100% (pct). 5000 = 100% в OOXML (1/50 процента)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")

    # 3) ширина рабочей области страницы (page_width - margins)
    section = doc.sections[0]
    total_width = section.page_width - section.left_margin - section.right_margin

    table.autofit = False

    # 4) распределение ширины по колонкам
    if cols <= 2:
        left_ratio = max(0.05, min(0.95, float(left_ratio)))
        w0 = int(total_width * left_ratio)
        w1 = int(total_width - w0)
        widths = [w0, w1]
    else:
        w = int(total_width / cols)
        widths = [w] * cols

    # 5) применяем ширины к колонкам и существующим ячейкам
    for i in range(min(cols, len(table.columns))):
        table.columns[i].width = widths[i]

    for row in table.rows:
        for i in range(min(cols, len(row.cells))):
            row.cells[i].width = widths[i]


def strip_parentheses(text: str | None) -> str:
    if not text:
        return "-"
    return text.split("(", 1)[0].strip()


def set_cell_text(cell, text, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    set_run_font(run, bold=bold)


def fill_table(table, obj: dict, sections, json_formatter):
    for section_title, fields in sections:
        rows = []
        for field, label in fields:
            raw = obj.get(field)

            if field.endswith("_json"):
                value = json_formatter(raw)
            else:
                value = format_value(field, raw)

            # Не выводим пустое. "-" не выводим (если нужно выводить — убери "-")
            if value is None:
                continue
            if isinstance(value, str) and value.strip() in ("", "-"):
                continue

            rows.append((label, value))

        if not rows:
            continue

        add_section_header_row(table, section_title)
        for label, value in rows:
            r = table.add_row().cells
            set_cell_text(r[0], label)
            set_cell_text(r[1], value)


def fmt(value):
    """Для таблицы зон поражающих факторов:
    - None -> '-'
    - если |value| >= 1: округлить до целого (int)
    - если |value| < 1: округлить до 1 знака после запятой
    """
    if value is None:
        return "-"

    try:
        v = float(value)
    except Exception:
        return str(value)

    if abs(v) >= 1:
        return str(int(round(v)))
    return f"{v:.1f}"


def render_section_at_marker(
        doc: Document,
        marker: str,
        section_title: str,
        items: list[dict],
        item_title_field: str,
        sections,
        json_formatter,
):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(doc, p_marker, section_title)
    if p.runs:
        set_run_font(p.runs[0], bold=True)

    cur = p
    for item in items:
        item_title = item.get(item_title_field, "—")
        cur = insert_paragraph_after(doc, cur, str(item_title))
        if cur.runs:
            set_run_font(cur.runs[0], bold=True)

        table = insert_table_after(doc, cur, rows=0, cols=2, style="Table Grid")
        fill_table(table, item, sections, json_formatter)

        cur = insert_paragraph_after_table(doc, table, "")


def render_substances_one_table_at_marker(
        doc: Document,
        marker: str,
        items: list[dict],
        item_title_field: str,
        sections,
        json_formatter,
):
    """
    Одна таблица на все вещества без лишних абзацев:
    - Шапка "Параметр / Значение" (1 раз)
    - Разделитель между веществами: строка с названием вещества (merge 2 ячейки)
    - Параграф с маркером удаляем, чтобы не было пустой строки перед таблицей
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    # очищаем маркер (на всякий случай)
    clear_paragraph(p_marker)

    # вставляем таблицу сразу "после маркера"
    table = insert_table_after(doc, p_marker, rows=1, cols=2, style="Table Grid")

    # удаляем параграф-маркер, чтобы не было пустой строки между "Таблица 1 ..." и таблицей
    delete_paragraph(p_marker)
    # растянуть таблицу по ширине
    set_table_full_width(doc, table, cols=2)

    # шапка
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Параметр", bold=True)
    set_cell_text(hdr[1], "Значение", bold=True)

    for item in items or []:
        # разделитель вещества
        item_title = item.get(item_title_field, "—")
        sep = table.add_row().cells
        merged = sep[0].merge(sep[1])
        set_cell_text(merged, str(item_title), bold=True)

        # наполнение (те же секции/поля)
        fill_table(table, item, sections, json_formatter)


def render_equipment_one_table_at_marker(
        doc: Document,
        marker: str,
        items: list[dict],
        item_title_field: str,
        sections,
        json_formatter,
):
    """
    Одна таблица на всё оборудование без лишних абзацев:
    - "Параметр / Значение" 1 раз
    - разделитель между единицами: строка с названием оборудования (merge 2 cells)
    - жирные заголовки секций ОСТАЮТСЯ
    - исключаем внутренние параметры расчета по LABEL
    - параграф с маркером удаляем, чтобы не было пустой строки перед таблицей
    - таблицу растягиваем по ширине страницы
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    table = insert_table_after(doc, p_marker, rows=1, cols=2, style="Table Grid")

    # убираем пустую строку (маркерный абзац)
    delete_paragraph(p_marker)

    # растянуть таблицу по ширине
    set_table_autofit_to_contents(table)

    # шапка
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Параметр", bold=True)
    set_cell_text(hdr[1], "Значение", bold=True)

    excluded_labels = {
        "Тип оборудования",
        "Тип координат",
        "Координаты",
        "Поражение персонала",
        "Возможные погибшие, чел.",
        "Возможные пострадавшие, чел.",
    }

    filtered_sections = []
    for section_title_i, fields in sections:
        new_fields = [(field, label) for (field, label) in fields if label not in excluded_labels]
        filtered_sections.append((section_title_i, new_fields))

    for item in items or []:
        item_title = item.get(item_title_field, "—")

        # разделитель оборудования
        sep = table.add_row().cells
        merged = sep[0].merge(sep[1])
        set_cell_text(merged, str(item_title), bold=True)

        # заполняем (заголовки секций остаются)
        fill_table(table, item, filtered_sections, json_formatter)


def render_distribution_table_at_marker(
        doc: Document,
        marker: str,
        title: str,
        rows: list[dict],
        *,
        equipment_items: list[dict] | None = None,
):
    """
    DISTRIBUTION_SECTION (многоуровневая шапка):
    1) "Технологический блок, оборудование" -> 4 подколонки:
       - Наименование составляющей (hazard_component из equipment)
       - Оборудование (equipment_name)
       - Вещество (substance_name)
       - Кол-во, ед (quantity_equipment из equipment; для трубопроводов всегда 1)
    2) "Количество опасного вещества, т" -> 2:
       - В единице оборудования (amount_t)
       - В блоке:
           * equipment_type != 0: amount_t * quantity_equipment
           * equipment_type == 0: amount_t
    3) "Физические условия содержания опасного вещества" -> 3:
       - Агр. состояние (phase_state)
       - Давление, МПа (pressure_mpa)
       - Температура, °C (substance_temperature_c)

    Шапка 2 строки, обе повторяемые.
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # 9 колонок: 4 + 2 + 3
    table = insert_table_after(doc, p_marker, rows=2, cols=9, style="Table Grid")
    delete_paragraph(p_marker)

    # растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=9, left_ratio=1 / 9)

    # ---- шапка (2 строки) ----
    r0 = table.rows[0].cells
    r1 = table.rows[1].cells

    # Верхняя строка: 3 "больших" объединённых заголовка
    set_cell_text(r0[0], "Технологический блок, оборудование", bold=True)
    r0[0].merge(r0[3])
    set_cell_text(r0[4], "Количество опасного вещества, т", bold=True)
    r0[4].merge(r0[5])
    set_cell_text(r0[6], "Физические условия содержания опасного вещества", bold=True)
    r0[6].merge(r0[8])

    # Нижняя строка: подзаголовки
    headers2 = [
        "Наименование составляющей",
        "Оборудование",
        "Вещество",
        "Кол-во, ед",
        "В единице оборудования",
        "В блоке",
        "Агр. состояние",
        "Давление, МПа",
        "Температура, °C",
    ]
    for i, h in enumerate(headers2):
        set_cell_text(r1[i], h, bold=True)

    # повторяемые строки шапки
    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])

    # ---- мапы из оборудования: component + quantity + type ----
    eq_component: dict[str, object] = {}
    eq_quantity: dict[str, object] = {}
    eq_type: dict[str, object] = {}

    for e in (equipment_items or []):
        name = e.get("equipment_name")
        if not name:
            continue
        name = str(name)
        eq_component[name] = e.get("hazard_component")
        eq_quantity[name] = e.get("quantity_equipment")
        eq_type[name] = e.get("equipment_type")

    # ---- строки ----
    for r in rows or []:
        eq_name = r.get("equipment_name") or "-"
        eq_key = str(eq_name)

        substance = strip_parentheses(r.get("substance_name")) or "-"
        phase = r.get("phase_state") if r.get("phase_state") is not None else "-"

        comp = eq_component.get(eq_key, "-")
        equipment_type = eq_type.get(eq_key)

        # amount_t
        amount = r.get("amount_t")
        try:
            amount_num = float(amount) if amount is not None else None
        except Exception:
            amount_num = None

        # qty + amount_block (по ТЗ)
        if equipment_type == 0:
            # трубопровод
            qty_num = 1
            amount_block = amount_num
        else:
            qty = eq_quantity.get(eq_key)
            try:
                qty_num = float(qty) if qty is not None else None
            except Exception:
                qty_num = None

            amount_block = (
                amount_num * qty_num
                if amount_num is not None and qty_num is not None
                else None
            )

        row = table.add_row().cells
        set_cell_text(row[0], comp if comp is not None else "-")
        set_cell_text(row[1], eq_name)
        set_cell_text(row[2], substance)

        # Кол-во, ед
        if isinstance(qty_num, (int, float)) and float(qty_num).is_integer():
            qty_out = int(qty_num)
        else:
            qty_out = qty_num if qty_num is not None else "-"
        set_cell_text(row[3], qty_out)

        # Количество ОВ
        set_cell_text(row[4], format_float_3(amount_num))
        set_cell_text(row[5], format_float_3(amount_block))

        # Физические условия
        set_cell_text(row[6], phase)
        set_cell_text(row[7], format_float_2(r.get("pressure_mpa")))
        set_cell_text(row[8], format_float_1(r.get("substance_temperature_c")))


def render_ov_amount_table_at_marker(doc: Document, marker: str, title: str, rows: list[dict]):
    """
    Таблица: Оценка количества опасного вещества в аварии (без лишних абзацев).
    Колонки:
      1) Наименование оборудования
      2) Номер сценария (С{scenario_no})
      3) Количество ОВ участвующего в аварии, т (ov_in_accident_t)
      4) Количество ОВ в создании поражающего фактора, т (ov_in_hazard_factor_t)
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    table = insert_table_after(doc, p_marker, rows=1, cols=4, style="Table Grid")
    delete_paragraph(p_marker)

    # растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=4, left_ratio=1 / 5)

    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Наименование оборудования", bold=True)
    set_cell_text(hdr[1], "Номер сценария", bold=True)
    set_cell_text(hdr[2], "Количество ОВ участвующего в аварии, т", bold=True)
    set_cell_text(hdr[3], "Количество ОВ в создании поражающего фактора, т", bold=True)

    # Порядок С1..Сn, внутри номера — по оборудованию
    rows_sorted = sorted(
        rows,
        key=lambda r: (
            int(r.get("scenario_no")) if r.get("scenario_no") is not None else 10 ** 9,
            str(r.get("equipment_name") or ""),
        ),
    )

    for idx, r in enumerate(rows_sorted, start=1):
        row = table.add_row().cells
        set_cell_text(row[0], idx)
        set_cell_text(row[1], r.get("equipment_name") or "-")
        sc_no = r.get("scenario_no")
        set_cell_text(row[0], r.get("equipment_name") or "-")
        set_cell_text(row[1], f"С{sc_no}" if sc_no is not None else "-")
        set_cell_text(row[2], format_float_3(r.get("ov_in_accident_t")))
        set_cell_text(row[3], format_float_3(r.get("ov_in_hazard_factor_t")))

    # пустой абзац после таблицы НЕ добавляем


def render_personnel_casualties_table_at_marker(doc: Document, marker: str, title: str, rows: list[dict]):
    """
    CASUALTIES_SECTION:
    - без "№ п/п"
    - добавлена колонка "Потерпевшие" = fatalities_count + injured_count
    - без лишних абзацев (удаляем маркерный параграф)
    - на всю ширину окна/страницы
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # Колонки (примерно как было, но без "№ п/п" + новая "Потерпевшие"):
    # 0 Оборудование
    # 1 Номер сценария
    # 2 Погибшие
    # 3 Пострадавшие
    # 4 Потерпевшие (=2+3)
    table = insert_table_after(doc, p_marker, rows=1, cols=5, style="Table Grid")
    delete_paragraph(p_marker)

    set_table_full_width(doc, table, cols=5, left_ratio=1 / 5)

    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Наименование оборудования", bold=True)
    set_cell_text(hdr[1], "Номер сценария", bold=True)
    set_cell_text(hdr[2], "Погибшие, чел.", bold=True)
    set_cell_text(hdr[3], "Раненые, чел.", bold=True)
    set_cell_text(hdr[4], "Потерпевшие, чел.", bold=True)

    # Порядок: С1..Сn, внутри — по оборудованию (как делали в других таблицах)
    rows_sorted = sorted(
        rows,
        key=lambda r: (
            int(r.get("scenario_no")) if r.get("scenario_no") is not None else 10 ** 9,
            str(r.get("equipment_name") or ""),
        ),
    )

    for r in rows_sorted:
        eq_name = r.get("equipment_name") or "-"
        sc_no = r.get("scenario_no")
        sc_label = f"С{sc_no}" if sc_no is not None else "-"

        fat = r.get("fatalities_count")
        inj = r.get("injured_count")

        # безопасное суммирование
        try:
            fat_n = int(fat) if fat is not None else 0
        except Exception:
            fat_n = 0
        try:
            inj_n = int(inj) if inj is not None else 0
        except Exception:
            inj_n = 0

        victims = fat_n + inj_n

        row = table.add_row().cells
        set_cell_text(row[0], eq_name)
        set_cell_text(row[1], sc_label)
        set_cell_text(row[2], fat_n if fat is not None else "-")
        set_cell_text(row[3], inj_n if inj is not None else "-")
        set_cell_text(row[4], victims)


def _load_typical_scenarios() -> dict:
    """typical_scenarios.json не изменяем; читаем из data/typical_scenarios.json."""
    p = TYPICAL_SCENARIOS_PATH
    if p.exists():
        with open(p, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def _scenario_item_to_text(item) -> str:
    if isinstance(item, str):
        return item
    if isinstance(item, dict):
        # поддержка {"scenario_text": "..."} и близких вариантов
        return str(item.get("scenario_text") or item.get("text") or item.get("description") or "Описание не задано")
    return "Описание не задано"


def _get_scenarios_root(typical: dict) -> dict:
    # root может быть либо сразу словарём, либо лежать в ключе "scenarios"
    if isinstance(typical, dict) and "scenarios" in typical and isinstance(typical["scenarios"], dict):
        return typical["scenarios"]
    return typical if isinstance(typical, dict) else {}


def _get_description_list(root: dict, equipment_type, kind):
    # основной формат: root[equipment_type][kind] -> list
    et = str(equipment_type)
    kd = str(kind)
    if et in root and isinstance(root[et], dict) and kd in root[et]:
        return root[et][kd]
    # запасной вариант: root[kind][equipment_type]
    if kd in root and isinstance(root[kd], dict) and et in root[kd]:
        return root[kd][et]
    return None


def _build_scenario_index(rows: list[dict]):
    """
    Для каждой пары (equipment_type, kind) строим отображение:
    scenario_no (уникальный, отсортированный) -> локальный индекс 0..N-1
    """
    mapping = {}
    for r in rows:
        key = (r.get("equipment_type"), r.get("substance_kind"))
        mapping.setdefault(key, set()).add(r.get("scenario_no"))

    idx_map = {}
    for key, sc_set in mapping.items():
        sc_list = sorted([x for x in sc_set if x is not None])
        idx_map[key] = {sc_no: i for i, sc_no in enumerate(sc_list)}
    return idx_map


def render_scenarios_table_at_marker(doc: Document, marker: str, title: str, rows: list[dict]):
    """
    Таблица сценариев (без лишних абзацев):

    1) Наименование оборудования
    2) Номер сценария (С{scenario_no})
    3) Описание сценария
    4) Базовая частота, 1/год
    5) Условная вероятность, -
    6) Частота сценария аварии, 1/год
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # Вставляем таблицу сразу после маркера и удаляем маркер, чтобы не было пустой строки
    table = insert_table_after(doc, p_marker, rows=1, cols=6, style="Table Grid")
    delete_paragraph(p_marker)

    # Растянуть по ширине окна/страницы
    set_table_autofit_to_contents(table)

    hdr = table.rows[0].cells
    headers = [
        "Наименование оборудования",
        "Номер сценария",
        "Описание сценария",
        "Базовая частота, 1/год",
        "Условная вероятность, -",
        "Частота сценария аварии, 1/год",
    ]
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)

    typical = _load_typical_scenarios()
    root = _get_scenarios_root(typical)

    # Порядок: С1..Сn, внутри — по оборудованию
    rows_sorted = sorted(
        rows,
        key=lambda r: (
            int(r.get("scenario_no")) if r.get("scenario_no") is not None else 10 ** 9,
            str(r.get("equipment_name") or ""),
        ),
    )

    for i, r in enumerate(rows_sorted, start=1):
        et = r.get("equipment_type")
        kd = r.get("substance_kind")
        sc_no = r.get("scenario_no")

        local_idx = r.get("scenario_idx")
        desc_list = _get_description_list(root, et, kd)
        if isinstance(desc_list, list) and local_idx is not None and 0 <= local_idx < len(desc_list):
            desc = _scenario_item_to_text(desc_list[local_idx])
        else:
            desc = "Описание не задано"

        row = table.add_row().cells
        set_cell_text(row[0], r.get("equipment_name") or "-")
        set_cell_text(row[1], f"С{sc_no}" if sc_no is not None else "-")
        set_cell_text(row[2], desc)
        set_cell_text(row[3], format_exp(r.get("base_frequency")))
        set_cell_text(row[4],
                      r.get("accident_event_probability") if r.get("accident_event_probability") is not None else "-")
        set_cell_text(row[5], format_exp(r.get("scenario_frequency")))
        set_table_full_width(doc, table, cols=6)


def render_impact_zones_table(doc: Document, marker: str, rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # Вставляем таблицу сразу после маркера
    table = insert_table_after(doc, p_marker, rows=1, cols=22, style="Table Grid")

    # Удаляем маркерный абзац -> нет лишней строки между "Таблица ..." и таблицей
    delete_paragraph(p_marker)

    # Растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=22, left_ratio=1 / 22)

    headers = [
        "Наименование оборудования",
        "Номер сценария",
        "q = 10,5",
        "q = 7,0",
        "q = 4,2",
        "q = 1,4",
        "ΔР = 70",
        "ΔР = 28",
        "ΔР = 14",
        "ΔР = 5",
        "ΔР = 2",
        "Lf",
        "Df",
        "Rнкпр",
        "Rвсп",
        "Rlpt",
        "Rppt",
        "Q600",
        "Q320",
        "Q220",
        "Q120",
        "St",
    ]

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for idx, r in enumerate(rows, start=1):
        row = table.add_row().cells

        set_cell_text(row[0], r.get("equipment_name") or "-")
        set_cell_text(row[1], f"С{r.get('scenario_no')}" if r.get("scenario_no") is not None else "-")

        set_cell_text(row[2], fmt(r.get("q_10_5")))
        set_cell_text(row[3], fmt(r.get("q_7_0")))
        set_cell_text(row[4], fmt(r.get("q_4_2")))
        set_cell_text(row[5], fmt(r.get("q_1_4")))

        set_cell_text(row[6], fmt(r.get("p_70")))
        set_cell_text(row[7], fmt(r.get("p_28")))
        set_cell_text(row[8], fmt(r.get("p_14")))
        set_cell_text(row[9], fmt(r.get("p_5")))
        set_cell_text(row[10], fmt(r.get("p_2")))

        set_cell_text(row[11], fmt(r.get("l_f")))
        set_cell_text(row[12], fmt(r.get("d_f")))
        set_cell_text(row[13], fmt(r.get("r_nkpr")))
        set_cell_text(row[14], fmt(r.get("r_vsp")))
        set_cell_text(row[15], fmt(r.get("l_pt")))
        set_cell_text(row[16], fmt(r.get("p_pt")))

        set_cell_text(row[17], fmt(r.get("q_600")))
        set_cell_text(row[18], fmt(r.get("q_320")))
        set_cell_text(row[19], fmt(r.get("q_220")))
        set_cell_text(row[20], fmt(r.get("q_120")))

        set_cell_text(row[21], fmt(r.get("s_t")))


def render_damage_table_at_marker(doc, marker: str, rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # без лишних абзацев: таблица сразу на месте маркера
    table = insert_table_after(doc, p_marker, rows=1, cols=8, style="Table Grid")
    delete_paragraph(p_marker)

    # растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=8, left_ratio=1 / 8)

    headers = [
        "Наименование оборудования",
        "Номер сценария",
        "Прямые потери, тыс.руб",
        "Затраты на ликвидацию, тыс.руб",
        "Социальные потери, тыс.руб",
        "Косвенный ущерб, тыс.руб",
        "Экологический ущерб, тыс.руб",
        "Суммарный ущерб, тыс.руб",
    ]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("equipment_name", "-"))
        set_cell_text(row[1], f"С{r.get('scenario_no')}" if r.get("scenario_no") is not None else "-")

        set_cell_text(row[2], format_float_1(r.get("direct_losses")))
        set_cell_text(row[3], format_float_1(r.get("liquidation_costs")))
        set_cell_text(row[4], format_float_1(r.get("social_losses")))
        set_cell_text(row[5], format_float_1(r.get("indirect_damage")))
        set_cell_text(row[6], format_float_1(r.get("total_environmental_damage")))
        set_cell_text(row[7], format_float_1(r.get("total_damage")))

    # пустой абзац после таблицы НЕ добавляем
    # insert_paragraph_after_table(doc, table, "")


def render_collective_risk_table(doc, marker: str, rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # таблица сразу на месте маркера
    table = insert_table_after(doc, p_marker, rows=1, cols=3, style="Table Grid")
    delete_paragraph(p_marker)

    # растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=3, left_ratio=1 / 3)

    headers = [
        "Составляющая ОПО",
        "Коллективный риск гибели, чел·год⁻¹",
        "Коллективный риск ранения, чел·год⁻¹",
    ]

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("hazard_component", "-"))
        set_cell_text(row[1], format_exp(r.get("collective_risk_fatalities")))
        set_cell_text(row[2], format_exp(r.get("collective_risk_injured")))


def render_individual_risk_table(doc, marker: str, rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # таблица сразу на месте маркера
    table = insert_table_after(doc, p_marker, rows=1, cols=3, style="Table Grid")
    delete_paragraph(p_marker)

    # растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=3, left_ratio=1 / 3)

    headers = [
        "Составляющая ОПО",
        "Индивидуальный риск гибели, 1·год⁻¹",
        "Индивидуальный риск ранения, 1·год⁻¹",
    ]

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("hazard_component", "-"))
        set_cell_text(row[1], format_exp(r.get("individual_risk_fatalities")))
        set_cell_text(row[2], format_exp(r.get("individual_risk_injured")))


def render_fatal_accident_frequency_text(doc, marker: str, min_freq, max_freq):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    if min_freq is None or max_freq is None:
        text = "Частота аварии с гибелью не менее одного человека равна 0."
    else:
        text = (
            "Частота аварии с гибелью не менее одного человека "
            f"лежит в диапазоне для объекта от {format_exp(min_freq)} "
            f"и до {format_exp(max_freq)} 1/год."
        )

    # Ничего не вставляем и не очищаем абзац целиком — сохраняем стиль из шаблона
    # Просто заменяем содержимое в первом run, остальные очищаем.
    if p_marker.runs:
        p_marker.runs[0].text = text
        for run in p_marker.runs[1:]:
            run.text = ""
    else:
        p_marker.add_run(text)


def render_max_damage_by_component_table(doc, marker: str, rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # Таблица сразу на месте маркера (без лишних абзацев/заголовков)
    table = insert_table_after(doc, p_marker, rows=1, cols=3, style="Table Grid")
    delete_paragraph(p_marker)

    # Растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=3, left_ratio=1 / 3)

    headers = [
        "Составляющая ОПО",
        "Максимальный суммарный ущерб, тыс.руб",
        "Максимальный экологический ущерб, тыс.руб",
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("hazard_component", "-"))
        set_cell_text(row[1], format_float_1(r.get("max_total_damage")))
        set_cell_text(row[2], format_float_1(r.get("max_total_environmental_damage")))

    # Пустой абзац после таблицы НЕ добавляем


from docx.shared import Cm
from pathlib import Path


def render_chart_at_marker(
        doc: Document,
        marker: str,
        title: str,
        image_path: Path | None,
        *,
        width_cm: float = 16.0,
        height_cm: float = 11.0,
):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    # Вставляем "на месте": ничего не добавляем после/до
    clear_paragraph(p_marker)

    # Если диаграммы нет — пишем текст в этот же абзац
    if image_path is None or not Path(image_path).exists():
        p_marker.add_run("Данные для построения диаграммы отсутствуют.")
        return

    # Картинку вставляем в этот же абзац
    run = p_marker.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def render_fn_chart_at_marker(doc: Document, marker: str, fn_rows: list[dict]):
    points = build_fn_points(fn_rows)
    if not points:
        image_path = None
    else:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "fn.png"
        save_fn_chart(points, image_path)

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="F/N - диаграмма",
        image_path=image_path,
        width_cm=16.0,
    )


def render_fg_chart_at_marker(doc: Document, marker: str, fg_rows: list[dict]):
    points = build_fg_points(fg_rows)
    if not points:
        image_path = None
    else:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "fg.png"
        save_fg_chart(points, image_path)

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="F/G - диаграмма",
        image_path=image_path,
        width_cm=16.0,
    )


def render_pareto_fatalities_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    """Pareto по коллективному риску гибели (collective_risk_fatalities)"""
    series = build_pareto_series(rows, "collective_risk_fatalities")
    charts_dir = OUT_PATH.parent / "charts"
    charts_dir.mkdir(parents=True, exist_ok=True)
    img_path = charts_dir / "pareto_fatalities.png" if series else None
    if series:
        save_pareto_chart(
            series,
            img_path,
            title="Pareto-диаграмма (вклад) сценариев по коллективному риску гибели",
            ylabel="Коллективный риск гибели, чел·год⁻¹",
        )
    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Pareto сценариев по коллективному риску гибели",
        image_path=img_path,
        width_cm=16.0,
    )


def render_pareto_injured_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    """Pareto по коллективному риску ранения (collective_risk_injured)"""
    series = build_pareto_series(rows, "collective_risk_injured")
    charts_dir = OUT_PATH.parent / "charts"
    charts_dir.mkdir(parents=True, exist_ok=True)
    img_path = charts_dir / "pareto_injured.png" if series else None
    if series:
        save_pareto_chart(
            series,
            img_path,
            title="Pareto-диаграмма (вклад) сценариев по коллективному риску ранения",
            ylabel="Коллективный риск ранения, чел·год⁻¹",
        )
    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Pareto сценариев по коллективному риску ранения",
        image_path=img_path,
        width_cm=16.0,
    )


def render_pareto_damage_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    # собираем серии: подпись = "equipment / Сn", значение = total_damage
    series = build_pareto_series(rows, value_key="total_damage")
    image_path = None

    if series:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "pareto_damage.png"

        save_pareto_chart(
            series=series,
            path=image_path,
            title="Pareto-диаграмма (вклад) сценариев по суммарному ущербу",
            ylabel="Суммарный ущерб, тыс.руб",
        )

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Pareto сценариев по суммарному ущербу",
        image_path=image_path,
        width_cm=16.0,
    )


def render_pareto_environmental_damage_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    series = build_pareto_series(rows, value_key="total_environmental_damage")
    image_path = None

    if series:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "pareto_environmental_damage.png"

        save_pareto_chart(
            series=series,
            path=image_path,
            title="Pareto-диаграмма (вклад) сценариев по экологическому ущербу",
            ylabel="Экологический ущерб, тыс.руб",
        )

    render_chart_at_marker(
        doc,
        marker,
        "Pareto сценариев по экологическому ущербу",
        image_path,
    )


def render_component_damage_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    image_path = None
    if rows:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "damage_by_component.png"
        save_component_damage_chart(rows, image_path)

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Распределение ущерба по составляющим ОПО",
        image_path=image_path,
        width_cm=16.0,
    )


def render_risk_matrix_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    image_path = None
    if rows:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "risk_matrix.png"
        save_risk_matrix_chart(rows, image_path)

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Матрица риска (частота – последствия)",
        image_path=image_path,
        width_cm=16.0,
    )


def render_risk_matrix_damage_chart_at_marker(doc: Document, marker: str, rows: list[dict]):
    image_path = None
    if rows:
        charts_dir = OUT_PATH.parent / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        image_path = charts_dir / "risk_matrix_damage.png"
        save_risk_matrix_chart_damage(rows, image_path)

    render_chart_at_marker(
        doc=doc,
        marker=marker,
        title="Матрица риска (частота – ущерб)",
        image_path=image_path,
        width_cm=16.0,
    )


def render_top_scenarios_by_component_table(doc, marker: str, rows: list[dict]):
    """
    Колонки:
      1) Составляющая объекта (hazard_component)
      2) Тип сценария (наиболее опасный / наиболее вероятный)
      3) Номер сценария (С{scenario_no})
      4) Оборудование (equipment_name)
      5) Количество погибших, чел (fatalities_count)
      6) Ущерб, тыс.руб (total_damage)
      7) Частота сценария, 1/год (scenario_frequency)
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # таблица сразу на месте маркера
    table = insert_table_after(doc, p_marker, rows=1, cols=7, style="Table Grid")

    # убрать лишний абзац между "Таблица ..." и таблицей
    delete_paragraph(p_marker)

    # растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=7, left_ratio=1 / 7)

    headers = [
        "Составляющая объекта",
        "Тип сценария",
        "Номер сценария",
        "Оборудование",
        "Количество погибших, чел",
        "Ущерб, тыс.руб",
        "Частота сценария, 1/год",
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    type_map = {
        "dangerous": "Наиболее опасный",
        "probable": "Наиболее вероятный",
    }

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("hazard_component", "-"))
        set_cell_text(row[1], type_map.get(r.get("scenario_type"), str(r.get("scenario_type", "-"))))

        sc_no = r.get("scenario_no")
        set_cell_text(row[2], f"С{sc_no}" if sc_no is not None else "-")

        set_cell_text(row[3], r.get("equipment_name", "-"))
        set_cell_text(row[4], r.get("fatalities_count", 0))
        set_cell_text(row[5], format_float_1(r.get("total_damage")))
        set_cell_text(row[6], format_exp(r.get("scenario_frequency")))

    # пустой абзац после таблицы НЕ добавляем
    # insert_paragraph_after_table(doc, table, "")


def render_fatality_risk_by_component_table(doc, marker: str, rows: list[dict]):
    """Сводная таблица: индивидуальный и коллективный риск гибели по составляющим ОПО."""
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # таблица сразу на месте маркера (без лишних абзацев)
    table = insert_table_after(doc, p_marker, rows=1, cols=3, style="Table Grid")
    delete_paragraph(p_marker)

    # растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=3, left_ratio=1 / 3)

    headers = [
        "Составляющая ОПО",
        "Индивидуальный риск гибели, 1·год⁻¹",
        "Коллективный риск гибели, чел·год⁻¹",
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for r in rows:
        row = table.add_row().cells
        set_cell_text(row[0], r.get("hazard_component", "-"))
        set_cell_text(row[1], format_exp(r.get("individual_risk_fatalities")))
        set_cell_text(row[2], format_exp(r.get("collective_risk_fatalities")))

    # пустой абзац после таблицы НЕ добавляем


def render_comparative_fatality_risk_table(doc, marker: str, individual_risk_rows: list[dict]):
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # таблица сразу на месте маркера (без лишних абзацев)
    table = insert_table_after(doc, p_marker, rows=1, cols=2, style="Table Grid")
    delete_paragraph(p_marker)

    # растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=2)

    # Заголовки
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Вид смертельной опасности", bold=True)
    set_cell_text(hdr[1], "Уровень риска, дБR", bold=True)

    # 1) Социальные риски
    for name, dbr in SOCIAL_FATALITY_RISKS_DBR:
        row = table.add_row().cells
        set_cell_text(row[0], str(name))
        set_cell_text(row[1], f"{dbr:+.1f}")

    # 2) ОПО по составляющим (берем уже посчитанный индивидуальный риск)
    for r in individual_risk_rows:
        comp = r.get("hazard_component")
        risk = (
            r.get("individual_risk_fatalities")
            if r.get("individual_risk_fatalities") is not None
            else r.get("individual_risk")
        )

        dbr = risk_to_dbr(risk)
        dbr_txt = "—" if dbr is None else f"{dbr:+.1f}"

        row = table.add_row().cells
        set_cell_text(row[0], f"Риск гибели при аварии на ОПО ({comp})")
        set_cell_text(row[1], dbr_txt)

    # пустой абзац после таблицы НЕ добавляем


def render_ngk_background_comparison_table(doc, marker: str, conn):
    p = find_paragraph_with_marker(doc, marker)
    if p is None:
        return

    clear_paragraph(p)

    # Таблица сразу на месте маркера (без title-абзаца)
    table = insert_table_after(doc, p, rows=1, cols=2, style="Table Grid")

    # Удаляем маркерный абзац, чтобы не было пустой строки между "Таблица ..." и таблицей
    delete_paragraph(p)

    # Растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=2, left_ratio=0.55)  # можно 0.5, но 55/45 обычно читабельнее

    # Заголовки
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Параметр", bold=True)
    set_cell_text(hdr[1], "Значение", bold=True)

    # --- A. Фон НГК ---
    for name, value in NGK_BACKGROUND_RISK:
        row = table.add_row().cells
        set_cell_text(row[0], name)
        set_cell_text(row[1], value)

    # --- B. ОПО по составляющим ---
    risks = get_individual_risk(conn)
    damage_by_comp = get_damage_by_component(conn)

    for r in risks:
        comp = r["hazard_component"]
        risk = r["individual_risk_fatalities"]

        if risk is None or risk <= 0:
            continue

        dbr = risk_to_dbr(risk)
        ppm = risk * 1e6
        damage = damage_by_comp.get(comp)

        # Ущерб
        row = table.add_row().cells
        set_cell_text(row[0], f"Ущерб при аварии на ОПО ({comp}), тыс. руб")
        set_cell_text(row[1], f"{damage:.1f}" if damage is not None else "—")

        # Риск дБR
        row = table.add_row().cells
        set_cell_text(row[0], f"Риск гибели при аварии на ОПО ({comp}), дБR")
        set_cell_text(row[1], "—" if dbr is None else f"{dbr:+.1f}")

        # Риск ppm
        row = table.add_row().cells
        set_cell_text(row[0], f"Риск гибели при аварии на ОПО ({comp}), ppm")
        set_cell_text(row[1], f"{ppm:.2f}")


def render_substances_by_component_table(doc, marker: str, conn):
    """
    Таблица: "Составляющая объекта" / "Количество вещества, т"

    Расчет количества вещества должен соответствовать колонке "В блоке" из DISTRIBUTION:
      - equipment_type == 0 (трубопровод): amount_t (Кол-во, ед = 1)
      - equipment_type != 0: amount_t * quantity_equipment

    Источники:
      - get_used_equipment(conn) -> hazard_component, equipment_type, quantity_equipment
      - get_hazard_distribution(conn) -> equipment_name, substance_name, amount_t
    """
    p = find_paragraph_with_marker(doc, marker)
    if p is None:
        return

    clear_paragraph(p)

    # Таблица сразу на месте маркера (без заголовка-абзаца)
    table = insert_table_after(doc, p, rows=1, cols=2, style="Table Grid")
    delete_paragraph(p)

    # Растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=2, left_ratio=0.5)

    # Заголовки
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Составляющая объекта", bold=True)
    set_cell_text(hdr[1], "Количество вещества, т", bold=True)

    # --- Подтягиваем оборудование для мапов ---
    equipment = get_used_equipment(conn) or []

    eq_component: dict[str, str] = {}
    eq_quantity: dict[str, float] = {}
    eq_type: dict[str, int] = {}

    for e in equipment:
        name = e.get("equipment_name")
        if not name:
            continue
        key = str(name)

        eq_component[key] = e.get("hazard_component") or "-"
        eq_type[key] = e.get("equipment_type")

        q = e.get("quantity_equipment")
        try:
            eq_quantity[key] = float(q) if q is not None else 1.0
        except Exception:
            eq_quantity[key] = 1.0

    # --- Берем распределение (как в DISTRIBUTION_SECTION) ---
    dist_rows = get_hazard_distribution(conn) or []

    # Агрегация: comp -> substance -> mass
    agg: dict[str, dict[str, float]] = {}
    comp_order: list[str] = []
    sub_order: dict[str, list[str]] = {}

    for r in dist_rows:
        eq_name = r.get("equipment_name") or "-"
        eq_key = str(eq_name)

        comp = eq_component.get(eq_key, "-")
        et = eq_type.get(eq_key)  # может быть None
        qty = eq_quantity.get(eq_key, 1.0)

        # amount_t
        amount = r.get("amount_t")
        try:
            amount_num = float(amount) if amount is not None else None
        except Exception:
            amount_num = None
        if amount_num is None:
            continue

        # substance name (без скобок, как в других местах)
        substance = strip_parentheses(r.get("substance_name")) or "-"
        if substance == "-":
            continue

        # расчет "В блоке"
        if et == 0:
            amount_block = amount_num
        else:
            amount_block = amount_num * qty

        if comp not in agg:
            agg[comp] = {}
            comp_order.append(comp)
            sub_order[comp] = []

        if substance not in agg[comp]:
            agg[comp][substance] = 0.0
            sub_order[comp].append(substance)

        agg[comp][substance] += amount_block

    # --- Вывод таблицы ---
    for comp in comp_order:
        items = agg.get(comp) or {}

        row = table.add_row().cells
        set_cell_text(row[0], str(comp))

        # во второй ячейке — список веществ построчно
        cell = row[1]
        cell.text = ""
        p_cell = cell.paragraphs[0]

        for i, name in enumerate(sub_order.get(comp, [])):
            mass = items.get(name)
            if mass is None:
                continue
            if i > 0:
                p_cell.add_run("\n")
            p_cell.add_run(f"{name} — {mass:.3f} т")


def render_top_scenarios_description_by_component_table(doc: Document, marker: str, conn):
    """
    Таблица:
    - Составляющая объекта
    - Тип сценария (наиболее опасный / наиболее вероятный)
    - Номер сценария (Сn)
    - Наименование оборудования
    - Краткое описание сценария
    - Частота, 1/год

    ВАЖНО: Описание берём НЕ из top_rows (там часто нет scenario_idx),
    а из get_scenarios(conn) по ключу (equipment_name, scenario_no).
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # таблица сразу на месте маркера
    table = insert_table_after(doc, p_marker, rows=1, cols=6, style="Table Grid")
    delete_paragraph(p_marker)

    # растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=6, left_ratio=1 / 6)

    headers = [
        "Составляющая объекта",
        "Тип сценария",
        "Номер сценария",
        "Оборудование",
        "Краткое описание сценария",
        "Частота, 1/год",
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    def _scenario_type_label(st: str | None) -> str:
        if st == "dangerous":
            return "Наиболее опасный"
        if st == "probable":
            return "Наиболее вероятный"
        return str(st or "-")

    # 1) Топ-сценарии
    top_rows = get_top_scenarios_by_hazard_component(conn) or []

    # 2) Карта сценариев (equipment_name, scenario_no) -> {equipment_type, substance_kind, scenario_idx}
    scenario_rows = get_scenarios(conn) or []
    sc_map: dict[tuple[str, int], dict] = {}
    for rr in scenario_rows:
        eqn = rr.get("equipment_name")
        scn = rr.get("scenario_no")
        if eqn is None or scn is None:
            continue
        try:
            scn_i = int(scn)
        except Exception:
            continue
        sc_map[(str(eqn), scn_i)] = rr

    # 3) Типовые описания
    typical = _load_typical_scenarios()
    root = _get_scenarios_root(typical)

    for r in top_rows:
        comp = r.get("hazard_component")
        st = r.get("scenario_type")
        sc_no = r.get("scenario_no")
        eq_name = r.get("equipment_name")
        freq = r.get("scenario_frequency")

        # --- Описание сценария ---
        desc = "Описание не задано"

        eq_key = str(eq_name) if eq_name is not None else None
        try:
            sc_i = int(sc_no) if sc_no is not None else None
        except Exception:
            sc_i = None

        base = sc_map.get((eq_key, sc_i)) if (eq_key is not None and sc_i is not None) else None
        if base:
            et = base.get("equipment_type")
            kd = base.get("substance_kind")
            local_idx = base.get("scenario_idx")  # 0..N-1

            desc_list = _get_description_list(root, et, kd)
            if isinstance(desc_list, list) and local_idx is not None:
                try:
                    li = int(local_idx)
                except Exception:
                    li = None
                if li is not None and 0 <= li < len(desc_list):
                    desc = _scenario_item_to_text(desc_list[li])

        row = table.add_row().cells
        set_cell_text(row[0], comp if comp is not None else "-")
        set_cell_text(row[1], _scenario_type_label(st))
        set_cell_text(row[2], f"С{sc_no}" if sc_no is not None else "-")
        set_cell_text(row[3], eq_name or "-")
        set_cell_text(row[4], desc)
        set_cell_text(row[5], format_exp(freq))

    # пустой абзац после таблицы не добавляем


def render_top_scenarios_pf_by_component_table(doc: Document, marker: str, conn):
    """
    Таблица:
    - Составляющая объекта
    - Тип сценария
    - Номер сценария (Сn)
    - Наименование оборудования
    - Зоны действия поражающих факторов
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # ✅ таблица сразу на месте маркера (без заголовка-абзаца)
    table = insert_table_after(doc, p_marker, rows=1, cols=5, style="Table Grid")

    # ✅ убрать пустую строку между "Таблица ..." и таблицей
    delete_paragraph(p_marker)

    # ✅ растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=5, left_ratio=1 / 5)

    hdr = table.rows[0].cells
    headers = [
        "Составляющая объекта",
        "Тип сценария",
        "Номер сценария",
        "Наименование оборудования",
        "Зоны действия поражающих факторов",
    ]
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)

    top_rows = get_top_scenarios_by_hazard_component(conn)

    def _type_label(st):
        return "Наиболее опасный" if st == "dangerous" else "Наиболее вероятный"

    for r in top_rows:
        comp = r.get("hazard_component")
        st = r.get("scenario_type")
        sc_no = r.get("scenario_no")
        eq_name = r.get("equipment_name")

        calc_row = get_calculation_row_for_top_scenario(conn, comp, sc_no, eq_name)
        zones_txt = _format_pf_zones(calc_row)

        row = table.add_row().cells
        set_cell_text(row[0], comp if comp is not None else "-")
        set_cell_text(row[1], _type_label(st))
        set_cell_text(row[2], f"С{sc_no}" if sc_no is not None else "-")
        set_cell_text(row[3], eq_name or "Поражающие факторы отсуствуют")
        set_cell_text(row[4], zones_txt)

    # ❌ пустой абзац после таблицы НЕ добавляем
    # insert_paragraph_after_table(doc, table, "")


def render_top_scenarios_fatalities_injured_by_component_table(doc: Document, marker: str, conn):
    """
    Таблица:
    - Составляющая объекта
    - Тип сценария
    - Номер сценария (Сn)
    - Наименование оборудования
    - Количество погибших и пострадавших
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    table = insert_table_after(doc, p_marker, rows=1, cols=5, style="Table Grid")
    delete_paragraph(p_marker)
    hdr = table.rows[0].cells

    headers = [
        "Составляющая объекта",
        "Тип сценария",
        "Номер сценария",
        "Наименование оборудования",
        "Количество погибших, раненых и пострадавших",
    ]
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)

    top_rows = get_top_scenarios_by_hazard_component(conn)

    def _type_label(st):
        return "Наиболее опасный" if st == "dangerous" else "Наиболее вероятный"

    for r in top_rows:
        comp = r.get("hazard_component")
        st = r.get("scenario_type")
        sc_no = r.get("scenario_no")
        eq_name = r.get("equipment_name")

        fi = get_fatalities_injured_for_top_scenario(conn, comp, sc_no, eq_name)
        if fi is None:
            text = "Погибшие: 0; Пострадавшие: 0"
        else:
            fat, inj = fi
            fat_txt = str(int(fat)) if fat is not None else "0"
            inj_txt = str(int(inj)) if inj is not None else "0"
            text = f"Погибшие: {fat_txt}; Раненые: {inj_txt}; Пострадавшие: {str(int(inj_txt) + int(fat_txt))}"

        row = table.add_row().cells
        set_cell_text(row[0], comp if comp is not None else "-")
        set_cell_text(row[1], _type_label(st))
        set_cell_text(row[2], f"С{sc_no}" if sc_no is not None else "-")
        set_cell_text(row[3], eq_name or "-")
        set_cell_text(row[4], text)


def render_top_scenarios_damage_by_component_table(doc: Document, marker: str, conn):
    """
    Таблица:
    - Составляющая объекта
    - Тип сценария
    - Номер сценария (Сn)
    - Наименование оборудования
    - Ущерб, тыс.руб

    Формат:
    - без лишнего абзаца между "Таблица X" и таблицей (таблица на месте маркера)
    - растягивание таблицы по ширине окна/страницы
    - без пустого абзаца после таблицы
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # Таблица сразу на месте маркера
    table = insert_table_after(doc, p_marker, rows=1, cols=5, style="Table Grid")

    # Удаляем маркерный абзац, чтобы не было пустой строки перед таблицей
    delete_paragraph(p_marker)

    # Растянуть по ширине окна/страницы
    set_table_full_width(doc, table, cols=5, left_ratio=0.22)

    hdr = table.rows[0].cells
    headers = [
        "Составляющая объекта",
        "Тип сценария",
        "Номер сценария",
        "Наименование оборудования",
        "Ущерб, тыс.руб",
    ]
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)

    top_rows = get_top_scenarios_by_hazard_component(conn)

    def _type_label(st):
        return "Наиболее опасный" if st == "dangerous" else "Наиболее вероятный"

    for r in top_rows:
        comp = r.get("hazard_component")
        st = r.get("scenario_type")
        sc_no = r.get("scenario_no")
        eq_name = r.get("equipment_name")

        dmg = get_total_damage_for_top_scenario(conn, comp, sc_no, eq_name)
        dmg_txt = f"{float(dmg):.1f}" if dmg is not None else "—"

        row = table.add_row().cells
        set_cell_text(row[0], comp if comp is not None else "-")
        set_cell_text(row[1], _type_label(st))
        set_cell_text(row[2], f"С{sc_no}" if sc_no is not None else "-")
        set_cell_text(row[3], eq_name or "-")
        set_cell_text(row[4], dmg_txt)

    # пустой абзац после таблицы НЕ добавляем
    # insert_paragraph_after_table(doc, table, "")


def render_top_scenarios_final_conclusion_table(doc: Document, marker: str, conn):
    """
    Заключительная таблица по наиболее опасному/вероятному сценарию по каждой составляющей.
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    p = insert_paragraph_after(
        doc, p_marker,
        "Заключительная таблица по наиболее опасному и наиболее вероятному сценарию"
    )
    if p.runs:
        set_run_font(p.runs[0], bold=True)

    table = insert_table_after(doc, p, rows=1, cols=11, style="Table Grid")
    hdr = table.rows[0].cells

    headers = [
        "Составляющая объекта",
        "Тип сценария",
        "Номер сценария",
        "Наименование оборудования",
        "Краткое описание сценария",
        "Частота, 1/год",
        "Количество опасного вещества участвующего в аварии, т",
        "Зоны действия поражающих факторов",
        "Метод расчёта",
        "Количество погибших и пострадавших",
        "Ущерб, тыс.руб",
    ]
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)

    # 1) Топ-сценарии
    top_rows = get_top_scenarios_by_hazard_component(conn)

    # 2) Для описания сценария — используем (equipment_name, scenario_no) + scenario_idx
    scenario_rows = get_scenarios(conn)
    typical = _load_typical_scenarios()
    root = _get_scenarios_root(typical)

    sc_map = {}
    for rr in scenario_rows:
        key = (rr.get("equipment_name"), rr.get("scenario_no"))
        if key[0] is None or key[1] is None:
            continue
        sc_map.setdefault(key, rr)

    def _type_label(st: str) -> str:
        return "Наиболее опасный" if st == "dangerous" else "Наиболее вероятный"

    for r in top_rows:
        comp = r.get("hazard_component")
        st = r.get("scenario_type")
        sc_no = r.get("scenario_no")
        eq_name = r.get("equipment_name")
        freq = r.get("scenario_frequency")

        # --- Описание сценария ---
        desc = "Описание не задано"
        base = sc_map.get((eq_name, sc_no))
        if base:
            et = base.get("equipment_type")
            kd = base.get("substance_kind")
            local_idx = base.get("scenario_idx")  # 0..N-1 внутри оборудования

            desc_list = _get_description_list(root, et, kd)
            if isinstance(desc_list, list) and local_idx is not None and 0 <= local_idx < len(desc_list):
                desc = _scenario_item_to_text(desc_list[local_idx])

        # --- Поражающие факторы + методика ---
        calc_pf = get_calculation_row_for_top_scenario(conn, comp, sc_no, eq_name)
        zones_txt = _format_pf_zones(calc_pf)
        method_txt = _detect_method_text(calc_pf)

        # --- OV in accident ---
        ov = get_ov_in_accident_for_top_scenario(conn, comp, sc_no, eq_name)
        ov_txt = f"{float(ov):.2f}" if ov is not None else "—"

        # --- Погибшие/пострадавшие ---
        fi = get_fatalities_injured_for_top_scenario(conn, comp, sc_no, eq_name)
        if fi is None:
            fi_txt = "Погибшие: —; Пострадавшие: —"
        else:
            fat, inj = fi
            fat_txt = str(int(fat)) if fat is not None else "—"
            inj_txt = str(int(inj)) if inj is not None else "—"
            fi_txt = f"Погибшие: {fat_txt}; Пострадавшие: {inj_txt}"

        # --- Ущерб ---
        dmg = get_total_damage_for_top_scenario(conn, comp, sc_no, eq_name)
        dmg_txt = f"{float(dmg):.1f}" if dmg is not None else "—"

        row = table.add_row().cells
        set_cell_text(row[0], comp if comp is not None else "-")
        set_cell_text(row[1], _type_label(st))
        set_cell_text(row[2], f"С{sc_no}" if sc_no is not None else "-")
        set_cell_text(row[3], eq_name or "-")
        set_cell_text(row[4], desc)
        set_cell_text(row[5], format_exp(freq))
        set_cell_text(row[6], ov_txt)
        set_cell_text(row[7], zones_txt)
        set_cell_text(row[8], method_txt)
        set_cell_text(row[9], fi_txt)
        set_cell_text(row[10], dmg_txt)

    insert_paragraph_after_table(doc, table, "")


def _safe_json(v):
    if isinstance(v, dict):
        return v
    if isinstance(v, str):
        v = v.strip()
        if v.startswith("{") and v.endswith("}"):
            try:
                return json.loads(v)
            except Exception:
                return None
    return None


def render_substances_info_table_at_marker(doc: Document, marker: str, conn):
    """
    Таблица "Сведения об опасных веществах"
    Колонки:
      1) Наименование опасного вещества (без скобок)
      2) Степень опасности и характер воздействия (класс опасности, reactivity, impact, first_aid)

    Вещества: ВСЕ вещества из БД (не только используемые в расчётах).
    """
    p_marker = find_paragraph_with_marker(doc, marker)
    if p_marker is None:
        return

    clear_paragraph(p_marker)

    # Вставляем таблицу прямо после маркера (без дополнительного заголовка/абзаца)
    table = insert_table_after(doc, p_marker, rows=1, cols=2, style="Table Grid")
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Наименование опасного вещества", bold=True)
    set_cell_text(
        hdr[1],
        "Степень опасности и характер воздействия вещества на организм человека и характер воздействия веществ на окружающую среду",
        bold=True,
    )

    # --- ВСЕ вещества из БД ---
    try:
        from report.reportgen.db import get_all_substances  # надо добавить/иметь в db.py
    except Exception as e:
        raise RuntimeError(
            "Нужна функция get_all_substances(conn) в report.reportgen.db, "
            "чтобы вывести все вещества из БД (не только используемые)."
        ) from e

    all_substances = get_all_substances(conn) or []

    for s in all_substances:
        # 1) имя без скобок
        name = strip_parentheses(s.get("name"))

        # 2) класс опасности: поле или из toxicity_json/toxicity
        hazard_class = s.get("toxicity_hazard_class")

        if hazard_class is None:
            tox = _safe_json(s.get("toxicity_json"))
            if isinstance(tox, dict):
                hazard_class = tox.get("hazard_class")

        tox2 = _safe_json(s.get("toxicity"))
        if hazard_class is None and isinstance(tox2, dict):
            hazard_class = tox2.get("hazard_class")

        reactivity = s.get("reactivity")
        impact = s.get("impact")
        first_aid = s.get("first_aid")

        parts = [
            f"Класс опасности: {hazard_class if hazard_class is not None else '—'}",
            f"Реакционная способность: {reactivity if reactivity else '—'}",
            f"Воздействие: {impact if impact else '—'}",
            f"Первая помощь: {first_aid if first_aid else '—'}",
        ]

        row = table.add_row().cells
        set_cell_text(row[0], name)
        set_cell_text(row[1], "\n".join(parts))




def load_organization_root() -> dict:
    """Возвращает первый элемент массива из organization.json или {}."""
    p = ORGANIZATION_PATH
    if not p.exists():
        return {}
    with open(p, "r", encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, list) and data:
        return data[0]
    if isinstance(data, dict):
        return data
    return {}


def _norm_ws(s: str) -> str:
    # NBSP (U+00A0) и узкий NBSP (U+202F) -> обычный пробел
    return s.replace("\u00A0", " ").replace("\u202F", " ")


def _clear_run_direct_formatting(run):
    """
    Убирает прямое форматирование run (w:rPr), чтобы текст наследовал стиль абзаца/документа.
    """
    r = run._r
    rPr = r.find(qn("w:rPr"))
    if rPr is not None:
        r.remove(rPr)

def _replace_in_paragraph_runs(paragraph, replacements: dict) -> bool:
    changed = False

    # Для каких плейсхолдеров нужно сбрасывать прямой шрифт/размер
    RESET_KEYS = {"{{ NASF_INFORMATION }}", "{{ PASF_INFORMATION }}"}

    for run in paragraph.runs:
        txt = run.text
        if not txt:
            continue

        txt_norm = _norm_ws(txt)
        new_txt = txt_norm

        replaced_any = False
        replaced_reset_key = False

        for k, v in replacements.items():
            k_norm = _norm_ws(k)
            if k_norm in new_txt:
                new_txt = new_txt.replace(k_norm, v)
                replaced_any = True
                if k in RESET_KEYS:
                    replaced_reset_key = True

        if replaced_any:
            run.text = new_txt
            changed = True

            # ✅ сбросить прямое форматирование только для NASF/PASF
            if replaced_reset_key:
                _clear_run_direct_formatting(run)

        elif txt_norm != txt:
            # NBSP -> space
            run.text = txt_norm

    return changed



def _replace_in_paragraph_joined_runs(paragraph, replacements: dict) -> bool:
    if not paragraph.runs:
        return False

    full = "".join(run.text or "" for run in paragraph.runs)
    if not full:
        return False

    full = _norm_ws(full)
    new_full = full

    RESET_KEYS = {"{{ NASF_INFORMATION }}", "{{ PASF_INFORMATION }}"}
    saw_reset_key = False

    for k, v in replacements.items():
        k_norm = _norm_ws(k)
        if k_norm in new_full:
            new_full = new_full.replace(k_norm, v)
            if k in RESET_KEYS:
                saw_reset_key = True

    if new_full == full:
        return False

    paragraph.runs[0].text = new_full
    for run in paragraph.runs[1:]:
        run.text = ""

    # ✅ сбросить прямое форматирование, чтобы не “летел” шрифт
    if saw_reset_key:
        _clear_run_direct_formatting(paragraph.runs[0])

    return True

def calc_max_people_victims(casualties_rows: list[dict]) -> int:
    """
    MAX_PEOPLE_VICTIMS = max(fatalities_count + injured_count) по всем сценариям.
    None трактуем как 0.
    """
    max_v = 0
    for r in (casualties_rows or []):
        fat = r.get("fatalities_count")
        inj = r.get("injured_count")

        try:
            fat_n = int(fat) if fat is not None else 0
        except Exception:
            fat_n = 0

        try:
            inj_n = int(inj) if inj is not None else 0
        except Exception:
            inj_n = 0

        v = fat_n + inj_n
        if v > max_v:
            max_v = v
    return max_v


def _paragraph_has_any_placeholder(paragraph, replacements: dict) -> bool:
    full = "".join(run.text or "" for run in paragraph.runs)
    if not full:
        return False
    return any(k in full for k in replacements)  # тут можно передавать и dict, и dict_keys

# def replace_placeholders_in_paragraph(p, replacements):
#     # Получаем единый текст
#     text = p.text
#
#     # Делаем замены
#     for key, value in replacements.items():
#         variants = [
#             f"{{{{ {key} }}}}",
#             f"{{{{{key}}}}}",
#             f"{{{{ {key}}}}}",
#             f"{{{{{key} }}}}",
#         ]
#         for ph in variants:
#             text = text.replace(ph, str(value))
#
#     if text == p.text:
#         return  # ничего не изменилось
#
#     # Удаляем run’ы
#     for r in p.runs:
#         r._r.getparent().remove(r._r)
#
#     # Создаём один run
#     p.add_run(text)



def replace_placeholders_in_doc(doc: Document, replacements: dict):
    """
    Заменяет плейсхолдеры во всем документе: абзацы + таблицы (включая вложенные).
    """
    # 1) обычные абзацы
    for p in doc.paragraphs:
        _replace_in_paragraph_runs(p, replacements)
        if _paragraph_has_any_placeholder(p, replacements):
            _replace_in_paragraph_joined_runs(p, replacements)

    # 2) таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_runs(p, replacements)
                    if _paragraph_has_any_placeholder(p, replacements):
                        _replace_in_paragraph_joined_runs(p, replacements)


def select_site_by_id(sites: list[dict], site_id: str) -> dict:
    """
    Возвращает площадку с заданным site_id.
    Если не найдено — пустой dict.
    """
    for s in sites or []:
        if isinstance(s, dict) and s.get("site_id") == site_id:
            return s
    return {}


def build_org_replacements(org_root: dict) -> dict:
    org_root = org_root or {}

    org = org_root.get("organization", {}) or {}
    ids = org.get("ids", {}) or {}
    contacts = org.get("contacts", {}) or {}
    head = org.get("head", {}) or {}

    permits = org_root.get("permits", {}) or {}
    management_docs = org_root.get("management_docs", {}) or {}
    security = org_root.get("security_and_response", {}) or {}
    reserves = org_root.get("reserves", {}) or {}

    sites = org_root.get("sites", []) or []

    # <<< ВОТ ЗДЕСЬ ВЫБОР >>>
    site0 = select_site_by_id(sites, ORGANIZATION_SITE_ID)
    personnel = site0.get("personnel", {}) or {}

    def s(x):
        return "" if x is None else str(x)

    return {
        # --- organization ---
        "{{ FULL_NAME }}": s(org.get("full_name")),
        "{{ SHORT_NAME }}": s(org.get("short_name")),

        "{{ OGRN }}": s(ids.get("ogrn")),
        "{{ INN }}": s(ids.get("inn")),
        "{{ KPP }}": s(ids.get("kpp")),

        "{{ ORG_ADDRESS }}": s(org.get("address")),

        "{{ ORG_EMAIL }}": s(contacts.get("email")),
        "{{ ORG_PHONE }}": s(contacts.get("phone")),
        "{{ ORG_FAX }}": s(contacts.get("fax")),

        "{{ HEAD_POSITION }}": s(head.get("position")),
        "{{ HEAD_FULL_NAME }}": s(head.get("full_name")),
        "{{ HEAD_SHORT_NAME }}": s(head.get("short_name")),

        # --- permits ---
        "{{ LICENSE_NUMBER }}": s(permits.get("license_number")),

        # --- management_docs ---
        "{{ INDUSTRIAL_SAFETY_MANAGEMENT_SYSTEM }}": s(management_docs.get("industrial_safety_management_system")),
        "{{ INDUSTRIAL_CONTROL_REGULATION }}": s(management_docs.get("industrial_control_regulation")),
        "{{ ACCIDENT_INVESTIGATION_REGULATION }}": s(management_docs.get("accident_investigation_regulation")),

        # --- security_and_response ---
        "{{ OPO_SECURITY }}": s(security.get("opo_security")),
        "{{ NASF_INFORMATION }}": s(security.get("nasf_information")),
        "{{ PASF_INFORMATION }}": s(security.get("pasf_information")),

        # --- reserves ---
        "{{ FINANCIAL_RESERVE_ORDER }}": s(reserves.get("financial_reserve_order")),
        "{{ MATERIAL_RESERVE_ORDER }}": s(reserves.get("material_reserve_order")),

        # --- site[0] ---
        "{{ SITE_ID }}": s(site0.get("site_id")),
        "{{ SITE_NAME }}": s(site0.get("name")),
        "{{ SITE_REG_NUMBER }}": s(site0.get("reg_number")),
        "{{ SITE_OBJECT_ID }}": s(site0.get("object_id")),

        "{{ SITE_OBJECT_ADDRESS }}": s(site0.get("object_address")),
        "{{ SITE_SANITARY_PROTECTION_ZONE_M }}": s(site0.get("sanitary_protection_zone_m")),

        "{{ SITE_DESCRIPTION }}": s(site0.get("description")),
        "{{ SITE_AREA_CHARACTERISTICS }}": s(site0.get("area_characteristics")),

        "{{ SITE_EMPLOYEES_COUNT }}": s(personnel.get("employees_count")),
        "{{ SITE_EMPLOYEES_OTHER_OPO_COUNT }}": s(personnel.get("employees_other_opo_count")),

        "{{ SITE_EMERGENCY_RESPONSE_PLAN }}": s(site0.get("emergency_response_plan")),
    }


def build_project_common_replacements(common: dict) -> dict:
    common = common or {}
    ex = common.get("executor", {}) or {}

    def s(x):
        return "" if x is None else str(x)

    return {
        # --- project ---
        "{{ PROJECT_YEAR }}": s(common.get("year")),
        "{{ PROJECT_NAME }}": s(common.get("project_name")),
        "{{ PROJECT_CODE }}": s(common.get("project_code")),
        "{{ DPB_CODE }}": s(common.get("dpb_code")),
        "{{ GOCHS_CODE }}": s(common.get("gochs_code")),
        "{{ PB_CODE }}": s(common.get("pb_code")),

        # --- executor ---
        "{{ EXECUTOR_NAME }}": s(ex.get("name")),
        "{{ EXECUTOR_ADDRESS }}": s(ex.get("address")),
        "{{ EXECUTOR_SRO }}": s(ex.get("sro")),
        "{{ EXECUTOR_INN }}": s(ex.get("inn")),
        "{{ EXECUTOR_OGRN }}": s(ex.get("ogrn")),
        "{{ EXECUTOR_TEL }}": s(ex.get("tel")),
        "{{ EXECUTOR_EMAIL }}": s(ex.get("email")),
        "{{ EXECUTOR_WEBSITE }}": s(ex.get("website")),
        "{{ EXECUTOR_HEAD_POSITION }}": s(ex.get("head_position")),
        "{{ EXECUTOR_HEAD_FULL_NAME }}": s(ex.get("head_full_name")),
        "{{ EXECUTOR_SPECIALIST_INFO }}": s(ex.get("specialist_info")),
    }

def fill_headers_footers(doc, replacements):
    def process_hf(hf):
        if not hf:
            return
        # абзацы колонтитула
        for p in hf.paragraphs:
            _replace_in_paragraph_runs(p, replacements)
            if _paragraph_has_any_placeholder(p, replacements):
                _replace_in_paragraph_joined_runs(p, replacements)

        # таблицы внутри колонтитула
        for tbl in hf.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph_runs(p, replacements)
                        if _paragraph_has_any_placeholder(p, replacements):
                            _replace_in_paragraph_joined_runs(p, replacements)

    for section in doc.sections:
        process_hf(section.header)
        process_hf(section.first_page_header)   # первая страница
        process_hf(section.even_page_header)
        process_hf(section.footer)
        process_hf(section.first_page_footer)   # первая страница футер
        process_hf(section.even_page_footer)


def fill_doc(
        doc: Document,
        *,
        substances,
        equipment,
        distribution,
        scenarios,
        ov_amounts,
        impact_zones,
        casualties,
        damage_rows,
        collective_risk_rows,
        individual_risk_rows,
        min_f,
        max_f,
        max_damage_rows,
        top_scenarios_rows,
        fatality_risk_by_component_rows,
        conn,
        fn_rows,
        fg_rows,
        pareto_rows,
        pareto_damage_rows,
        pareto_env_rows,
        component_damage_rows,
        risk_matrix_rows,
        risk_matrix_damage_rows,
):
    # Текстовые данные
    org_root = load_organization_root()
    proj_common = load_project_common()
    repl = build_org_replacements(org_root)
    repl.update(build_project_common_replacements(proj_common))
    # {{ MAX_PEOPLE_VICTIMS }} — максимальное число потерпевших (fatalities + injured)
    max_victims = calc_max_people_victims(casualties)
    repl["{{ MAX_PEOPLE_VICTIMS }}"] = str(max_victims)

    replace_placeholders_in_doc(doc, repl)

    # Внутри fill_doc после replace_placeholders_in_doc(doc, repl):
    fill_headers_footers(doc, repl)


    # Таблицы
    render_substances_one_table_at_marker(
        doc=doc,
        marker="{{SUBSTANCES_SECTION}}",
        items=substances,
        item_title_field="name",
        sections=SUBSTANCE_SECTIONS,
        json_formatter=pretty_json_substance,
    )

    render_substances_info_table_at_marker(
        doc=doc,
        marker="{{SUBSTANCES_INFO_SECTION}}",
        conn=conn,
    )

    render_equipment_one_table_at_marker(
        doc=doc,
        marker="{{EQUIPMENT_SECTION}}",
        items=equipment,
        item_title_field="equipment_name",
        sections=EQUIPMENT_SECTIONS,
        json_formatter=pretty_json_generic,
    )

    render_distribution_table_at_marker(
        doc=doc,
        marker="{{DISTRIBUTION_SECTION}}",
        title="Распределение опасного вещества по оборудованию",
        rows=distribution,
        equipment_items=equipment,
    )

    render_scenarios_table_at_marker(
        doc=doc,
        marker="{{SCENARIOS_SECTION}}",
        title="Сценарии аварий",
        rows=scenarios,
    )

    render_ov_amount_table_at_marker(
        doc=doc,
        marker="{{OV_AMOUNT_SECTION}}",
        title="Оценка количества опасного вещества в аварии",
        rows=ov_amounts,
    )

    render_impact_zones_table(doc, "{{IMPACT_ZONES_SECTION}}", impact_zones)

    render_personnel_casualties_table_at_marker(
        doc=doc,
        marker="{{CASUALTIES_SECTION}}",
        title="Оценка количества погибших/пострадавших",
        rows=casualties,
    )

    render_damage_table_at_marker(doc=doc, marker="{{DAMAGE_SECTION}}", rows=damage_rows)

    render_collective_risk_table(doc=doc, marker="{{COLLECTIVE_RISK_SECTION}}", rows=collective_risk_rows)
    render_individual_risk_table(doc=doc, marker="{{INDIVIDUAL_RISK_SECTION}}", rows=individual_risk_rows)

    render_fatal_accident_frequency_text(
        doc=doc,
        marker="{{FATAL_ACCIDENT_FREQUENCY}}",
        min_freq=min_f,
        max_freq=max_f,
    )

    render_max_damage_by_component_table(doc=doc, marker="{{MAX_DAMAGE_BY_COMPONENT_SECTION}}", rows=max_damage_rows)

    render_top_scenarios_by_component_table(
        doc=doc,
        marker="{{TOP_SCENARIOS_BY_COMPONENT_SECTION}}",
        rows=top_scenarios_rows,
    )

    render_fatality_risk_by_component_table(
        doc=doc,
        marker="{{FATALITY_RISK_BY_COMPONENT_SECTION}}",
        rows=fatality_risk_by_component_rows,
    )

    render_comparative_fatality_risk_table(
        doc=doc,
        marker="{{COMPARATIVE_FATALITY_RISK_TABLE}}",
        individual_risk_rows=individual_risk_rows,
    )

    # Эти функции используют conn
    render_ngk_background_comparison_table(doc=doc, marker="{{NGK_BACKGROUND_RISK_COMPARISON}}", conn=conn)
    render_substances_by_component_table(doc=doc, marker="{{SUBSTANCES_BY_COMPONENT_TABLE}}", conn=conn)

    render_top_scenarios_description_by_component_table(doc=doc, marker="{{TOP_SCENARIOS_DESC_BY_COMPONENT}}",
                                                        conn=conn)
    render_top_scenarios_pf_by_component_table(doc=doc, marker="{{TOP_SCENARIOS_PF_BY_COMPONENT}}", conn=conn)
    render_top_scenarios_fatalities_injured_by_component_table(
        doc=doc, marker="{{TOP_SCENARIOS_FATALITIES_INJURED}}", conn=conn
    )
    render_top_scenarios_damage_by_component_table(doc=doc, marker="{{TOP_SCENARIOS_DAMAGE}}", conn=conn)
    render_top_scenarios_final_conclusion_table(doc=doc, marker="{{TOP_SCENARIOS_FINAL_CONCLUSION}}", conn=conn)

    # Диаграммы (используют OUT_PATH.parent/"charts")
    render_fn_chart_at_marker(doc, "{{FN_CHART}}", fn_rows)
    render_fg_chart_at_marker(doc, "{{FG_CHART}}", fg_rows)
    render_pareto_damage_chart_at_marker(doc, "{{PARETO_DAMAGE_CHART}}", pareto_damage_rows)
    render_pareto_fatalities_chart_at_marker(doc, "{{PARETO_FATALITIES_CHART}}", pareto_rows)
    render_pareto_injured_chart_at_marker(doc, "{{PARETO_INJURED_CHART}}", pareto_rows)
    render_pareto_environmental_damage_chart_at_marker(doc, "{{PARETO_ENV_DAMAGE_CHART}}", pareto_env_rows)
    render_component_damage_chart_at_marker(doc, "{{DAMAGE_BY_COMPONENT_CHART}}", component_damage_rows)
    render_risk_matrix_chart_at_marker(doc, "{{RISK_MATRIX_CHART}}", risk_matrix_rows)
    render_risk_matrix_damage_chart_at_marker(doc, "{{RISK_MATRIX_DAMAGE_CHART}}", risk_matrix_damage_rows)


def main():
    REPORT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # 1) очищаем output от старых .docx
    clear_output_docx(REPORT_OUTPUT_DIR)

    # 2) берём все шаблоны текущего VARIANT
    templates = iter_variant_templates(TEMPLATE_DIR)
    # if not templates:
    #     # fallback: старое поведение
    #     templates = [REPORT_TEMPLATE_DOCX]

    with open_db(DB_PATH) as conn:
        # 3) собираем данные ОДИН РАЗ
        substances = get_all_substances(conn)
        equipment = get_used_equipment(conn)
        distribution = get_hazard_distribution(conn)
        scenarios = get_scenarios(conn)
        ov_amounts = get_ov_amounts_in_accident(conn)
        impact_zones = get_impact_zones(conn)
        casualties = get_personnel_casualties(conn)
        damage_rows = get_damage(conn)
        collective_risk_rows = get_collective_risk(conn)
        individual_risk_rows = get_individual_risk(conn)
        min_f, max_f = get_fatal_accident_frequency_range(conn)
        max_damage_rows = get_max_damage_by_hazard_component(conn)
        fn_rows = get_fn_source_rows(conn)
        fg_rows = get_fg_source_rows(conn)
        pareto_rows = get_pareto_risk_source_rows(conn)
        pareto_damage_rows = get_pareto_damage_source_rows(conn)
        pareto_env_rows = get_pareto_environmental_damage_source_rows(conn)
        component_damage_rows = get_max_losses_by_hazard_component(conn)
        risk_matrix_rows = get_risk_matrix_rows(conn)
        risk_matrix_damage_rows = get_risk_matrix_damage_rows(conn)
        top_scenarios_rows = get_top_scenarios_by_hazard_component(conn)

        # Сводная таблица рисков гибели по составляющим
        ind_map = {r.get("hazard_component"): r.get("individual_risk_fatalities") for r in individual_risk_rows}
        coll_map = {r.get("hazard_component"): r.get("collective_risk_fatalities") for r in collective_risk_rows}

        # ВАЖНО: без сортировок. Берём порядок как в исходных выборках из БД (первое появление).
        components = []
        seen = set()
        for src in (individual_risk_rows, collective_risk_rows):
            for rr in src:
                comp = rr.get("hazard_component")
                if comp is None or comp in seen:
                    continue
                seen.add(comp)
                components.append(comp)

        fatality_risk_by_component_rows = [
            {
                "hazard_component": comp,
                "individual_risk_fatalities": ind_map.get(comp),
                "collective_risk_fatalities": coll_map.get(comp),
            }
            for comp in components
        ]

        # 4) генерим все документы
        global OUT_PATH
        for template_path in templates:
            OUT_PATH = REPORT_OUTPUT_DIR / f"{template_path.stem}_out.docx"
            doc = Document(str(template_path))

            fill_doc(
                doc,
                substances=substances,
                equipment=equipment,
                distribution=distribution,
                scenarios=scenarios,
                ov_amounts=ov_amounts,
                impact_zones=impact_zones,
                casualties=casualties,
                damage_rows=damage_rows,
                collective_risk_rows=collective_risk_rows,
                individual_risk_rows=individual_risk_rows,
                min_f=min_f,
                max_f=max_f,
                max_damage_rows=max_damage_rows,
                top_scenarios_rows=top_scenarios_rows,
                fatality_risk_by_component_rows=fatality_risk_by_component_rows,
                conn=conn,
                fn_rows=fn_rows,
                fg_rows=fg_rows,
                pareto_rows=pareto_rows,
                pareto_damage_rows=pareto_damage_rows,
                pareto_env_rows=pareto_env_rows,
                component_damage_rows=component_damage_rows,
                risk_matrix_rows=risk_matrix_rows,
                risk_matrix_damage_rows=risk_matrix_damage_rows,
            )

            doc.save(str(OUT_PATH))
            print("Отчёт сформирован:", OUT_PATH)


if __name__ == "__main__":
    main()
